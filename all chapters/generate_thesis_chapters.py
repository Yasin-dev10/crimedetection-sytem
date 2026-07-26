"""
Generate JUST AI-FYP Chapters III, IV, and V for:
AUTOMATIC CLASSIFICATION OF CRIME-RELATED TEXT REPORTS USING NLP
(BAREAI system)
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

import json
import re

from thesis_expansions import inject as inject_expansions

ROOT = Path(__file__).resolve().parent
FIG = ROOT / "thesis_figures"
OUT = ROOT / "BAREAI_Complete_Thesis_Ch1-5.docx"
OUT_ALT = ROOT / "BAREAI_Complete_Thesis_FULL.docx"
OUT_CH345 = ROOT / "FYP_Chapters_3_4_5.docx"
OUT_CH123 = ROOT / "BAREAI_Chapters_1_2_3.docx"
CH12_JSON = ROOT / "_ch12_sections.json"

# Candidate details (IDs left blank for students to fill)
CANDIDATES = [
    ("YAASIIN MOHAMUUD ABDULLAAHI", "____________"),
    ("NAIMA ABDIAZIZ SAID", "____________"),
    ("NASTEHA MOHAMUUD MOHAMED", "____________"),
    ("NAJMA MUHIDIIN MOHAMED", "____________"),
]
PROJECT_TITLE = "AUTOMATIC CLASSIFICATION OF CRIME-RELATED TEXT REPORTS USING NLP"
DEGREE = "Bachelor of Science in Computer Applications"
FIELD = "Computer Applications"


def set_margins(section):
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(4.0)
    section.right_margin = Cm(2.0)


def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_paragraph(
    doc,
    text,
    bold=False,
    italic=False,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    space_before=0,
    space_after=0,
    line_spacing=2.0,
    first_line_indent=0,
    size=12,
):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    pf.first_line_indent = Cm(first_line_indent) if first_line_indent else Cm(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def configure_styles(doc):
    """Apply JUST formatting to Normal and Heading styles."""
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.left_indent = Cm(0)
    normal.paragraph_format.right_indent = Cm(0)

    for style_name, size, bold in [
        ("Heading 1", 12, True),
        ("Heading 2", 12, True),
        ("Heading 3", 12, True),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        pf = style.paragraph_format
        pf.space_before = Pt(12 if style_name == "Heading 1" else 10)
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 2.0
        pf.left_indent = Cm(0)
        pf.first_line_indent = Cm(0)
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_heading(doc, text, level=1):
    """Use real Word Heading styles so Table of Contents can update automatically."""
    style = f"Heading {min(level, 3)}"
    p = doc.add_paragraph(text, style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        set_run_font(run, size=12, bold=True, italic=(level >= 3))
    return p


def add_centered_title(doc, text, outline_level=None):
    """Centered section title for preliminary pages (kept in TOC via outline level)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(18)
    pf.line_spacing = 2.0
    if outline_level is not None:
        p.style = doc.styles["Heading 1"]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True)
    return p


def add_toc_entry(doc, title, page, level=0):
    """CONTENTS line with dot leaders (JUST/locust thesis style)."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.75 * level)
    # Right-aligned tab with dot leader near right margin (usable width ~14.5cm with 4cm left margin)
    tab_pos = Cm(14.0)
    p.paragraph_format.tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    run = p.add_run(title)
    set_run_font(run, size=12, bold=(level == 0 and title.isupper()))
    run2 = p.add_run(f"\t{page}")
    set_run_font(run2, size=12, bold=False)
    return p


def _add_page_number_field(paragraph, roman=False):
    """Insert PAGE field; optionally force Roman numeral format via section settings."""
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)
    set_run_font(run, size=8)


def _set_section_page_number(section, start=None, roman=False, show=True):
    section.footer.is_linked_to_previous = False
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # clear existing runs
    for r in list(p.runs):
        r._element.getparent().remove(r._element)

    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    if roman:
        pgNumType.set(qn("w:fmt"), "lowerRoman")
    else:
        pgNumType.set(qn("w:fmt"), "decimal")
    if start is not None:
        pgNumType.set(qn("w:start"), str(start))

    if show:
        _add_page_number_field(p, roman=roman)


def add_toc_field(paragraph):
    """Insert a Word TOC field (update with right-click → Update Field)."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    # placeholder text until Word updates the field
    text = OxmlElement("w:t")
    text.text = "Right-click this table of contents and choose Update Field."

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)
    set_run_font(run, size=12, italic=True)


def add_title_page(doc):
    for _ in range(2):
        add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    for line in [
        "AUTOMATIC CLASSIFICATION OF CRIME-RELATED",
        "TEXT REPORTS USING NLP",
    ]:
        add_paragraph(doc, line, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, space_after=4)
    for _ in range(2):
        add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    for name, _id in CANDIDATES:
        add_paragraph(doc, name, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, space_after=2)
    for _ in range(2):
        add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    for line in [
        "SUBMISSION OF GRADUATION PROJECT FOR",
        "PARTIAL FULFILLMENT OF THE",
        "DEGREE OF BACHELOR OF SCIENCE IN",
        "COMPUTER APPLICATIONS",
    ]:
        add_paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, space_after=2)
    for _ in range(2):
        add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    for line in [
        "JAMHURIYA UNIVERSITY OF SCIENCE AND",
        "TECHNOLOGY (JUST)",
        "FACULTY OF COMPUTER & INFORMATION",
        "TECHNOLOGY",
    ]:
        add_paragraph(doc, line, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, space_after=2)
    add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    add_paragraph(doc, "AUGUST 2026", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    # Title page has no printed page number; start prelim section next
    new_sec = doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(new_sec)
    _set_section_page_number(new_sec, start=1, roman=True, show=True)


def add_declaration(doc):
    """ORIGINALITY page — must fit on ONE page; signatures remain visible."""
    add_centered_title(doc, "ORIGINAL LITERARY WORK DECLARATION", outline_level=1)
    # Compact spacing so declaration + signatures stay on a single page
    for i, (name, cid) in enumerate(CANDIDATES, start=1):
        add_paragraph(
            doc,
            f"Name of Candidate {i}: {name}    ID: {cid}",
            bold=True,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            line_spacing=1.15,
            space_after=1,
            size=11,
        )
    add_paragraph(doc, f"Name of Degree: {DEGREE}", align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.15, space_after=1, size=11)
    add_paragraph(
        doc,
        'Title of Project/Thesis (“this Work”):',
        align=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.15,
        space_after=1,
        size=11,
    )
    add_paragraph(doc, PROJECT_TITLE, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.15, space_after=1, size=11)
    add_paragraph(doc, f"Field of Study: {FIELD}", align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.15, space_after=4, size=11)

    add_paragraph(
        doc,
        "We the undersigned do solemnly and sincerely declare that:",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        line_spacing=1.15,
        space_after=3,
        size=11,
    )
    declarations = [
        "(1) We are the authors/writers of this Work.",
        "(2) This Work is original.",
        "(3) Any use of any work in which copyright exists was done by way of fair dealing and for "
        "permitted purposes and any excerpt or extract from, or reference to or reproduction of any "
        "copyright work has been disclosed expressly and sufficiently and the title of the Work and "
        "its authorship have been acknowledged in this Work;",
        "(4) We do not have any actual knowledge nor ought we reasonably to know that the making of "
        "this work constitutes an infringement of any copyright work.",
        "(5) We hereby assign all and every right in the copyright to this Work to Jamhuriya "
        "University of Science and Technology (“JUST”), who henceforth shall be owner of the "
        "copyright in this Work and that any reproduction or use in any form or by any means "
        "whatsoever is prohibited without the written consent of JUST having been first had and obtained.",
        "(6) We are fully aware that if in the course of making this Work we have infringed any "
        "copyright whether intentionally or otherwise, we may be subject to legal action or any other "
        "action as may be determined by JUST.",
    ]
    for d in declarations:
        add_paragraph(doc, d, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.15, space_after=2, size=11)

    add_paragraph(doc, "", line_spacing=1.0, space_after=2)
    add_paragraph(
        doc,
        "Candidate 1’s Signature: ____________    Candidate 2’s Signature: ____________",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.15,
        space_after=2,
        size=11,
    )
    add_paragraph(
        doc,
        "Candidate 3’s Signature: ____________    Candidate 4’s Signature: ____________",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.15,
        space_after=2,
        size=11,
    )
    add_paragraph(doc, "Date: ____ / ____ / ________", align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.15, space_after=4, size=11)
    add_paragraph(
        doc,
        "Subscribed and solemnly declared before,",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.15,
        space_after=2,
        size=11,
    )
    add_paragraph(
        doc,
        "Supervisor’s Name: _______________________    Date: ____ / ____ / ________",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.15,
        space_after=2,
        size=11,
    )
    add_paragraph(
        doc,
        "Supervisor’s Signature: ___________________    Designation: ________________",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.15,
        space_after=2,
        size=11,
    )
    add_page_break(doc)


def add_dedication(doc):
    add_centered_title(doc, "DEDICATION", outline_level=1)
    add_paragraph(
        doc,
        "We dedicate our dissertation work to our family and many friends. A special feeling of "
        "gratitude goes to our loving parents, whose words of encouragement and push for tenacity "
        "ring in our ears. We also dedicate this work to the Somali community and to all who "
        "strive for safer digital spaces through education, technology, and public service.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        line_spacing=2.0,
        space_after=12,
    )
    add_page_break(doc)


def add_abstract(doc):
    add_centered_title(doc, "ABSTRACT", outline_level=1)
    abstract = (
        "Social media platforms continuously publish large volumes of unstructured text. In a "
        "security context, posts must be analysed and separated into crime-related and "
        "non-crime-related content so that investigators can prioritise genuine threats. Manual "
        "sorting of such posts is slow and error-prone. This study develops BAREAI, an Artificial "
        "Intelligence based Natural Language Processing (NLP) system that automatically classifies "
        "social media posts and related text reports as crime-related or non-crime-related. A "
        "labelled dataset of 9,999 Somali-oriented samples was used, comprising 5,000 crime-related "
        "and 4,999 non-crime-related texts. After preprocessing and TF-IDF feature extraction, "
        "several supervised models were trained and compared, including Logistic Regression, Support "
        "Vector Machine (SVM), and Random Forest. Random Forest achieved the best performance with "
        "approximately 89.55% accuracy, precision, recall, and F1-score on the held-out test set and "
        "was selected for deployment. The model is served through a Flask API and integrated into a "
        "web application that provides Text Analysis, File Analysis, URL Analysis, and Batch Analysis, "
        "together with case management and monitoring tools. The findings demonstrate that NLP-based "
        "classification can reduce manual triage effort and support faster crime detection from "
        "social media streams in low-resource settings."
    )
    add_paragraph(doc, abstract, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=2.0, space_after=12)
    add_paragraph(
        doc,
        "Keywords: Crime Classification, NLP, Machine Learning, Social Media, Crime Detection",
        bold=False,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        line_spacing=1.5,
        space_after=6,
    )
    add_page_break(doc)


def add_acknowledgements(doc):
    add_centered_title(doc, "ACKNOWLEDGEMENT", outline_level=1)
    paras = [
        "First and foremost, we thank Allah (SWT) for the guidance, strength, and opportunity to "
        "complete this Final Year Project successfully.",
        "Secondly, we express our deepest gratitude to our parents for their continuous moral and "
        "financial support throughout our academic journey.",
        "Thirdly, we thank our supervisor for valuable guidance, constructive feedback, and "
        "encouragement during the design, implementation, and writing of this thesis.",
        "Fourthly, we acknowledge Jamhuriya University of Science and Technology (JUST) for "
        "providing the academic environment and resources required to undertake this research.",
        "Finally, we thank our lecturers and classmates for knowledge shared, advice offered, and "
        "support given during the course of this project.",
    ]
    for para in paras:
        add_paragraph(doc, para, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=2.0, space_after=8)
    add_page_break(doc)


def add_contents(doc):
    """TABLE OF CONTENTS with dot leaders (JUST style)."""
    add_centered_title(doc, "TABLE OF CONTENTS", outline_level=1)

    # Preliminary pages (Roman numerals)
    prelim = [
        ("ORIGINAL LITERARY WORK DECLARATION", "i", 0),
        ("DEDICATION", "ii", 0),
        ("ABSTRACT", "iii", 0),
        ("ACKNOWLEDGEMENT", "iv", 0),
        ("TABLE OF CONTENTS", "v", 0),
        ("LIST OF FIGURES", "viii", 0),
        ("LIST OF TABLES", "x", 0),
        ("LIST OF ABBREVIATIONS", "xi", 0),
    ]
    for title, page, level in prelim:
        add_toc_entry(doc, title, page, level=level)

    # Main text
    main = [
        ("CHAPTER ONE: INTRODUCTION", "1", 0),
        ("1.1 Background of the Study", "1", 1),
        ("1.2 Problem Statement", "5", 1),
        ("1.3 Research Objectives", "7", 1),
        ("1.3.1 General Objective", "7", 2),
        ("1.3.2 Specific Objectives", "7", 2),
        ("1.4 Research Questions", "8", 1),
        ("1.5 Significance of the Study", "8", 1),
        ("1.6 Scope of the Study", "9", 1),
        ("1.7 Organization of the Study", "10", 1),
        ("CHAPTER TWO: LITERATURE REVIEW", "12", 0),
        ("2.1 Introduction", "12", 1),
        ("2.2 Impact of Crimes on Social Media", "16", 1),
        ("2.3 Traditional Classification Techniques in Crimes", "20", 1),
        ("2.4 Cybercrime Detection in Social Media", "24", 1),
        ("2.5 Challenges in the Classification of Crime-Related Text", "28", 1),
        ("2.6 Related Work", "32", 1),
        ("2.7 Proposed System", "36", 1),
        ("CHAPTER THREE: METHODOLOGY", "38", 0),
        ("3.1 Introduction", "38", 1),
        ("3.2 System Overview", "39", 1),
        ("3.3 System Features", "41", 1),
        ("3.4 System Architecture", "45", 1),
        ("3.5 Data Acquisition and Preprocessing", "48", 1),
        ("3.6 Feature Engineering / Representation", "51", 1),
        ("3.7 Model / Algorithm Development", "53", 1),
        ("3.8 Requirements", "57", 1),
        ("3.9 Feasibility Study", "60", 1),
        ("3.10 System Design", "62", 1),
        ("CHAPTER FOUR: IMPLEMENTATION AND RESULTS", "66", 0),
        ("4.1 Dataset Description", "66", 1),
        ("4.2 ML Pipeline Implementation", "68", 1),
        ("4.3 System Design and Deployment", "72", 1),
        ("4.4 Website Interface Screenshots", "76", 1),
        ("4.5 Data Visualization (Results Presentation)", "86", 1),
        ("4.6 Model Comparison", "90", 1),
        ("CHAPTER FIVE: DISCUSSION AND CONCLUSION", "93", 0),
        ("5.1 Introduction", "93", 1),
        ("5.2 Interpretation of Results", "94", 1),
        ("5.3 Contribution to Knowledge", "96", 1),
        ("5.4 Challenges and Limitations", "98", 1),
        ("5.5 Future Work", "99", 1),
        ("5.6 Final Conclusion", "100", 1),
        ("REFERENCES", "102", 0),
        ("APPENDIX A: FIGURE INDEX", "108", 0),
        ("APPENDIX B: TABLE INDEX", "110", 0),
    ]
    for title, page, level in main:
        add_toc_entry(doc, title, page, level=level)

    # Also embed an updatable Word TOC field for automatic page refresh
    add_paragraph(doc, "", line_spacing=1.0)
    note = add_paragraph(
        doc,
        "Automatic Table of Contents (update in Microsoft Word):",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.5,
        space_before=10,
        space_after=4,
    )
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    add_toc_field(p)
    tip = add_paragraph(
        doc,
        "Tip: In Microsoft Word, select the automatic TOC above → right-click → Update Field → Update entire table.",
        italic=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        line_spacing=1.5,
        space_before=4,
        space_after=6,
    )
    add_page_break(doc)


def add_list_of_figures(doc):
    add_centered_title(doc, "LIST OF FIGURES", outline_level=1)
    figures = [
        ("Figure 3.0: Multi-Source Input Analysis Features", "39"),
        ("Figure 3.1: BAREAI System Architecture Pipeline", "45"),
        ("Figure 3.2: Text Preprocessing Flowchart", "49"),
        ("Figure 3.3: NLP Crime Classification Pipeline", "54"),
        ("Figure 3.4: Use Case Diagram of BAREAI", "63"),
        ("Figure 3.5: Logical Database Schema of BAREAI", "64"),
        ("Figure 4.1: Dataset Class Distribution (N = 9,999)", "66"),
        ("Figure 4.2: End-to-End Classification and Case Dispatch Flow", "69"),
        ("Figure 4.3: ML Pipeline Implementation Stages", "70"),
        ("Figure 4.4: Local Deployment Topology of BAREAI Services", "74"),
        ("Figure 4.5: BAREAI Landing / Home Page", "76"),
        ("Figure 4.6: BAREAI Login Page", "77"),
        ("Figure 4.7: Text Analysis Page with Crime / Non-Crime Result", "78"),
        ("Figure 4.8: URL Analysis Page", "79"),
        ("Figure 4.9: File Analysis Page", "80"),
        ("Figure 4.10: Batch Analysis Page", "81"),
        ("Figure 4.11: Administrator Dashboard", "82"),
        ("Figure 4.12: Investigation Case Management Page", "83"),
        ("Figure 4.13: Blacklist and Monitoring Page", "84"),
        ("Figure 4.14: Notifications and Reports Page", "85"),
        ("Figure 4.15: Comparative Performance of Classification Models", "87"),
        ("Figure 4.16: Confusion Matrix of the Random Forest Classifier", "88"),
        ("Figure 4.17: Evaluation Metrics of the Selected Random Forest Model", "89"),
    ]
    for title, page in figures:
        add_toc_entry(doc, title, page, level=0)
    add_page_break(doc)


def add_list_of_tables(doc):
    add_centered_title(doc, "LIST OF TABLES", outline_level=1)
    tables = [
        ("Table 3.1: Dataset Column Description and Summary Statistics", "33"),
        ("Table 3.2: Feature Engineering and Split Configuration", "34"),
        ("Table 3.3: Model Evaluation Results on the Test Set", "35"),
        ("Table 3.4: Hardware Requirements", "37"),
        ("Table 3.5: Software Requirements", "38"),
        ("Table 3.6: Database Collections Summary", "41"),
        ("Table 4.1: Implemented Interface Modules", "46"),
        ("Table 4.2: Ranked Model Comparison Summary", "49"),
    ]
    for title, page in tables:
        add_toc_entry(doc, title, page, level=0)
    add_page_break(doc)


def add_list_of_abbreviations(doc):
    add_centered_title(doc, "LIST OF ABBREVIATIONS", outline_level=1)
    abbr = [
        ("AI", "Artificial Intelligence"),
        ("API", "Application Programming Interface"),
        ("FYP", "Final Year Project"),
        ("JSON", "JavaScript Object Notation"),
        ("JUST", "Jamhuriya University of Science and Technology"),
        ("ML", "Machine Learning"),
        ("NLP", "Natural Language Processing"),
        ("REST", "Representational State Transfer"),
        ("RF", "Random Forest"),
        ("SVM", "Support Vector Machine"),
        ("TF-IDF", "Term Frequency–Inverse Document Frequency"),
        ("URL", "Uniform Resource Locator"),
    ]
    for short, long in abbr:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(3.5), WD_TAB_ALIGNMENT.LEFT)
        run = p.add_run(f"{short}\t{long}")
        set_run_font(run, size=12)
    # body section starts next (no extra blank page break here)


def start_body_section(doc):
    """Start Arabic page numbering from 1 for Chapter I onward."""
    new_sec = doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(new_sec)
    _set_section_page_number(new_sec, start=1, roman=False, show=True)
    return new_sec


def add_body(doc, text, first_line_indent=1.27):
    """Body paragraph: Justify + first-line indent only (JUST thesis style)."""
    text = normalize_body_text(text)
    return add_paragraph(
        doc,
        text,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=first_line_indent,
        line_spacing=2.0,
        space_after=0,
        space_before=0,
    )


def normalize_body_text(text):
    """Clean PDF/extraction artifacts so lines do not look skewed."""
    if not text:
        return ""
    text = text.replace("\u00a0", " ").replace("\t", " ")
    text = text.replace("\r", " ").replace("\n", " ")
    text = re.sub(r" {2,}", " ", text)
    text = re.sub(r"\s+([,.;:?!)])", r"\1", text)
    text = re.sub(r"([(\[])\s+", r"\1", text)
    text = re.sub(r"\s+([)\]])", r"\1", text)
    # fix missing space after period before capital (common PDF join issue)
    text = re.sub(r"\.([A-Z])", r". \1", text)
    text = re.sub(r"\)\.([A-Z])", r"). \1", text)
    return text.strip()


def add_bullet(doc, text):
    text = normalize_body_text(text)
    p = add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(0)
    return p


def split_long_para(text, max_len=3500):
    """Keep paragraphs whole unless extremely long; never create fake mid-paragraph indents."""
    text = normalize_body_text(text)
    if len(text) <= max_len:
        return [text]
    sentences = re.split(r"(?<=[.!?])\s+(?=[A-Z(\"'])", text)
    chunks, cur = [], ""
    for s in sentences:
        if not s:
            continue
        if cur and len(cur) + len(s) + 1 > max_len:
            chunks.append(cur.strip())
            cur = s
        else:
            cur = (cur + " " + s).strip()
    if cur:
        chunks.append(cur.strip())
    return chunks or [text]


def add_table(doc, headers, rows, caption=None):
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        cap.paragraph_format.space_before = Pt(8)
        cap.paragraph_format.space_after = Pt(6)
        run = cap.add_run(caption)
        set_run_font(run, bold=True, size=11)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            for run in paragraph.runs:
                set_run_font(run, bold=True, size=10)

    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                for run in paragraph.runs:
                    set_run_font(run, size=10)

    doc.add_paragraph()
    return table


def add_figure(doc, image_path, caption, width=4.6):
    """Insert a centered figure with clear spacing so it does not collide with text."""
    path = Path(image_path)
    if not path.exists():
        alt = ROOT / path.name
        if alt.exists():
            path = alt
        else:
            add_paragraph(doc, f"[Missing figure: {path.name}]", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            add_paragraph(doc, caption, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0)
            return None

    add_page_break(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(min(width, 5.2)))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(18)
    run = cap.add_run(caption)
    set_run_font(run, bold=True, size=11)
    return cap


def add_page_break(doc):
    doc.add_page_break()


def add_chapters_one_two(doc):
    data = json.loads(CH12_JSON.read_text(encoding="utf-8"))
    sections = data["sections"]
    order = data["order"]

    heading_map = {
        "CHAPTER I: INTRODUCTION": (1, "CHAPTER ONE: INTRODUCTION"),
        "1.1 Background Of The Study": (2, "1.1 Background of the Study"),
        "1.2 Problem Statement": (2, "1.2 Problem Statement"),
        "1.3 Research Objectives": (2, "1.3 Research Objectives"),
        "1.3.1 General Objective": (3, "1.3.1 General Objective"),
        "1.3.2 Specific Objectives": (3, "1.3.2 Specific Objectives"),
        "1.4 Research Questions": (2, "1.4 Research Questions"),
        "1.5 Significance of the Study": (2, "1.5 Significance of the Study"),
        "1.6 Scope of the Study": (2, "1.6 Scope of the Study"),
        "1.7 Organization of the Study": (2, "1.7 Organization of the Study"),
        "CHAPTER II: LITERATURE REVIEW": (1, "CHAPTER TWO: LITERATURE REVIEW"),
        "2.1 Introduction": (2, "2.1 Introduction"),
        "2.2 Impact of crimes on social media": (2, "2.2 Impact of Crimes on Social Media"),
        "2.3 Traditional Classification techniques in crimes": (
            2,
            "2.3 Traditional Classification Techniques in Crimes",
        ),
        "2.4 Cybercrime detection in social media": (2, "2.4 Cybercrime Detection in Social Media"),
        "2.6 Related Work": (2, "2.6 Related Work"),
        "2.7 Proposed System": (2, "2.7 Proposed System"),
        "REFERENCES": (1, "REFERENCES"),
    }

    for key in order:
        if key == "REFERENCES":
            continue  # references placed at end of full thesis
        level_title = None
        for k, v in heading_map.items():
            if key.startswith(k) or key == k:
                level_title = v
                break
        if key.startswith("2.5"):
            level_title = (2, "2.5 Challenges in the Classification of Crime-Related Text")
        if level_title is None:
            level_title = (2, key)

        level, title = level_title
        if level == 1 and "CHAPTER TWO" in title:
            add_page_break(doc)
        add_heading(doc, title, level=level)

        # Clarifying opening for Chapter One after title
        if title == "CHAPTER ONE: INTRODUCTION":
            add_body(
                doc,
                "This chapter introduces the research problem and motivation for developing an "
                "automated NLP system that analyses social media posts and related digital text "
                "reports. Clearly, the proposed BAREAI system classifies each input as either "
                "crime-related or non-crime-related. The chapter presents the background, problem "
                "statement, objectives, research questions, significance, scope, and organisation "
                "of the study.",
            )
            add_body(
                doc,
                "In practical terms, BAREAI supports investigators and authorised users through "
                "four primary analysis features: Text Analysis (paste post content), File Analysis "
                "(PDF, DOCX, or TXT upload), URL Analysis (extract and classify webpage text), and "
                "Batch Analysis (classify multiple posts in one request). These features are "
                "detailed further in Chapter Three and demonstrated in Chapter Four.",
            )

        paras = sections.get(key, [])
        for para in paras:
            # bullets
            if para.strip().startswith("•") or para.strip().startswith("-"):
                for part in re.split(r"(?=•)", para):
                    part = part.strip()
                    if not part:
                        continue
                    add_bullet(doc, part)
                continue
            if re.match(r"^\d+\.\s", para.strip()) and "How to" in para:
                for part in re.split(r"(?=\d+\.\s)", para):
                    part = part.strip()
                    if part:
                        add_bullet(doc, part)
                continue
            for i, chunk in enumerate(split_long_para(para)):
                add_body(doc, chunk, first_line_indent=1.27 if i == 0 else 0)

        # Expansions for page-length and clarity
        if title == "1.7 Organization of the Study":
            inject_expansions(doc, "ch1_end", add_body)
        if title.startswith("2.") and title != "2.7 Proposed System":
            # short bridge paragraphs after each lit-review section
            add_body(
                doc,
                "The discussion above reinforces that social media text must be analysed "
                "automatically and separated into crime-related and non-crime-related classes. "
                "This insight directly informs the design choices presented in the methodology "
                "chapter, including Text, File, URL, and Batch analysis features.",
            )
        if title == "2.7 Proposed System":
            inject_expansions(doc, "ch2_extra", add_body)


def build():
    doc = Document()
    set_margins(doc.sections[0])
    configure_styles(doc)
    # Title page: no page number
    _set_section_page_number(doc.sections[0], start=1, roman=False, show=False)

    add_title_page(doc)  # creates prelim section with Roman numerals
    add_declaration(doc)
    add_dedication(doc)
    add_abstract(doc)
    add_acknowledgements(doc)
    add_contents(doc)
    add_list_of_figures(doc)
    add_list_of_tables(doc)
    add_list_of_abbreviations(doc)

    start_body_section(doc)  # Arabic numbering from 1
    add_chapters_one_two(doc)
    add_chapters_three_to_five(doc, page_break_before=True)
    save_thesis(doc, full=True)


def add_chapters_three_to_five(doc, page_break_before=True, through_chapter=5):
    """Add Chapter Three (and optionally Four–Five). through_chapter: 3 or 5."""
    # ==================================================================
    # CHAPTER III
    # ==================================================================
    if page_break_before:
        add_page_break(doc)
    add_heading(doc, "CHAPTER THREE: METHODOLOGY", level=1)

    add_heading(doc, "3.1 Introduction", level=2)
    add_body(
        doc,
        "This chapter presents the methodology used to design and develop BAREAI, an "
        "Artificial Intelligence based platform for the automatic classification of "
        "crime-related text reports using Natural Language Processing (NLP). The chapter "
        "describes the system overview, features, architecture, data acquisition and "
        "preprocessing procedures, feature engineering techniques, machine learning "
        "algorithms, evaluation metrics, system requirements, feasibility study, and "
        "system design. The methodology follows the Artificial Intelligence Final Year "
        "Project (FYP) guideline of the Faculty of Computer and Information Technology, "
        "Jamhuriya University of Science and Technology (JUST).",
    )
    add_body(
        doc,
        "The research problem addressed by this study is the heavy dependence of Somali "
        "security agencies on manual sorting of digital crime reports. Manual triage is "
        "slow, inconsistent, and difficult to scale when reports arrive from social media, "
        "websites, emails, and online forms. Therefore, the methodology combines classical "
        "supervised machine learning with a production-ready web application that supports "
        "investigators through automated classification, case dispatch, and monitoring.",
    )

    add_heading(doc, "3.2 System Overview", level=2)
    add_body(
        doc,
        "BAREAI is a web-based Artificial Intelligence platform that analyses social media "
        "posts and related digital text reports and classifies each input into one of two "
        "categories: crime-related or non-crime-related. The classification decision is the "
        "core output of the system and is returned together with a confidence score to support "
        "investigator triage.",
    )
    add_body(
        doc,
        "To make the system practical for real digital workflows, BAREAI provides four main "
        "analysis features: (1) Text Analysis — users paste a social media post or report "
        "directly into the interface; (2) File Analysis — users upload PDF, DOCX, or TXT "
        "documents; (3) URL Analysis — users provide a webpage link from which text is "
        "extracted and classified; and (4) Batch Analysis — users submit multiple posts in "
        "one request for bulk classification. These features reduce dependence on manual "
        "reading of every post.",
    )
    add_body(
        doc,
        "Beyond one-off analysis, the platform monitors blacklisted Facebook pages and "
        "websites, stores analysis history, creates investigation cases when crime-related "
        "content is detected, and notifies investigators. Role-based access control "
        "distinguishes Admin, Investigator, and Public User roles. Technically, BAREAI "
        "separates a React presentation layer, an Express application layer, a Flask NLP "
        "inference service hosting Random Forest with TF-IDF, and a MongoDB data layer.",
    )
    if (FIG / "fig3_6_input_features.png").exists():
        add_figure(
            doc,
            FIG / "fig3_6_input_features.png",
            "Figure 3.0: Multi-Source Input Analysis Features Leading to Binary Classification",
            width=4.7,
        )

    add_heading(doc, "3.3 System Features", level=2)

    add_heading(doc, "3.3.1 Authentication and Authorization", level=3)
    add_body(
        doc,
        "BAREAI implements secure registration and login with JSON Web Tokens (JWT). "
        "Email OTP verification strengthens account security. Investigators may be created "
        "by administrators with forced password change on first login. Roles restrict access "
        "to dashboards, case tools, blacklist management, and audit logs.",
    )

    add_heading(doc, "3.3.2 Automated Crime Text Classification", level=3)
    add_body(
        doc,
        "The core feature accepts Somali (and mixed) text reports and returns a binary "
        "prediction (crime-related / not crime-related), confidence score, optional matched "
        "keywords, and extracted location mentions. Classification is performed by the Flask "
        "NLP service and may be reinforced by Somali crime-keyword heuristics to reduce "
        "false negatives for high-risk phrases.",
    )

    add_heading(doc, "3.3.3 Multi-Source Input Analysis", level=3)
    add_body(
        doc,
        "Users can analyse: (a) pasted text; (b) webpage URLs extracted with HTML parsing; "
        "(c) uploaded documents (PDF, DOCX, TXT); and (d) batch submissions. Guest users are "
        "limited to a small number of free analyses to protect computational resources.",
    )

    add_heading(doc, "3.3.4 Blacklist Monitoring", level=3)
    add_body(
        doc,
        "Administrators and investigators maintain a watchlist of Facebook pages, websites, "
        "keywords, and persons. Facebook pages can be scanned periodically using automated "
        "browser scraping. Website pages can be scanned on demand. Detected crime content "
        "generates alerts and investigation cases.",
    )

    add_heading(doc, "3.3.5 Case Management and Notifications", level=3)
    add_body(
        doc,
        "When crime-related content is confirmed, the system creates a pending investigation "
        "case and notifies available investigators. The first investigator to accept the case "
        "claims ownership. Investigators can update status, write findings, submit formal "
        "reports, and flag false or malicious reports for administrative review.",
    )

    add_heading(doc, "3.3.6 Dashboard, Reports, and Audit Logging", level=3)
    add_body(
        doc,
        "The admin dashboard summarises analyses, detected crimes, users, investigators, "
        "blacklist items, and active cases. Operational reports and activity logs support "
        "accountability and organisational oversight.",
    )

    add_heading(doc, "3.4 System Architecture", level=2)
    add_body(
        doc,
        "The BAREAI architecture is expressed as an end-to-end pipeline rather than a narrow "
        "application-to-model link. Social media and related digital text enter the system "
        "through collection interfaces (text, file, URL, batch, and monitors). Collected text "
        "is preprocessed, converted into TF-IDF features, and classified by the machine "
        "learning model into crime-related or non-crime-related labels. Results are returned "
        "to the web application for display, history storage, and optional case dispatch.",
    )
    add_body(
        doc,
        "Figure 3.1 presents this pipeline from social media data to the web application. "
        "The flow emphasises that NLP inference is only one stage inside a complete "
        "operational system that also includes authentication, monitoring, and investigation "
        "support.",
    )
    add_figure(
        doc,
        FIG / "fig3_1_architecture.png",
        "Figure 3.1: BAREAI System Architecture Pipeline "
        "(Social Media Data to Web Application)",
        width=4.5,
    )

    add_heading(doc, "3.5 Data Acquisition and Preprocessing", level=2)
    add_heading(doc, "3.5.1 Dataset Source", level=3)
    add_body(
        doc,
        "The experimental dataset contains 9,999 labelled text samples stored in CSV format. "
        "Each record includes a source URL, the raw text body, and a binary category label: "
        "crime-related or not crime-related. The corpus is nearly balanced (5,000 "
        "crime-related and 4,999 not crime-related). Sources are predominantly Somali-language "
        "news and social-media style texts (for example BBC Somali and related outlets), "
        "manually labelled for supervised learning.",
    )
    add_table(
        doc,
        ["Attribute", "Description"],
        [
            ["url", "Source link of the collected text"],
            ["text", "Raw report / article / post content"],
            ["category", "Binary label: crime-related or not crime-related"],
            ["Total samples", "9,999"],
            ["Crime-related", "5,000 (50.01%)"],
            ["Not crime-related", "4,999 (49.99%)"],
            ["Average length", "Approximately 1,787 characters / 277 words"],
        ],
        caption="Table 3.1: Dataset Column Description and Summary Statistics",
    )

    add_heading(doc, "3.5.2 Data Preprocessing Steps", level=3)
    add_body(
        doc,
        "Preprocessing was implemented in a Jupyter notebook prior to model training. The "
        "pipeline removes incomplete rows, normalises casing, strips URLs and emails, retains "
        "Latin and Arabic script characters relevant to Somali orthography, collapses excess "
        "whitespace, tokenises by whitespace, and removes Somali stopwords (approximately "
        "284 stopwords when an external stopword list is available; otherwise a built-in "
        "fallback list is used).",
    )
    add_figure(
        doc,
        FIG / "fig3_5_preprocessing.png",
        "Figure 3.2: Text Preprocessing Flowchart",
        width=4.6,
    )

    add_heading(doc, "3.6 Feature Engineering / Representation", level=2)
    add_body(
        doc,
        "Text was transformed into numerical features using Term Frequency–Inverse Document "
        "Frequency (TF-IDF). During experimentation, TF-IDF and Bag-of-Words representations "
        "were compared with constrained vocabulary settings. For the notebook experiments, "
        "TF-IDF used max_features=300, min_df=2, max_df=0.95, and ngram_range=(1,2). The "
        "deployed production vectorizer uses a larger vocabulary (max_features=5000) to "
        "improve coverage during live inference.",
    )
    add_body(
        doc,
        "The dataset was split into training (80%) and testing (20%) subsets using stratified "
        "sampling with random_state=42, yielding 7,999 training samples and 2,000 test "
        "samples (1,000 per class in the test set). Stratification preserved class balance "
        "and enabled fair comparison across classifiers.",
    )
    add_table(
        doc,
        ["Item", "Configuration"],
        [
            ["Feature method", "TF-IDF (primary); Bag-of-Words (experimental)"],
            ["Notebook TF-IDF", "max_features=300, ngram_range=(1,2)"],
            ["Deployed TF-IDF", "max_features=5000, unigrams"],
            ["Train/Test split", "80% / 20%, stratified"],
            ["Random state", "42"],
            ["Target classes", "crime-related, not crime-related"],
        ],
        caption="Table 3.2: Feature Engineering and Split Configuration",
    )

    add_heading(doc, "3.7 Model / Algorithm Development", level=2)
    add_heading(doc, "3.7.1 Baseline Model and Model Selection", level=3)
    add_body(
        doc,
        "Multiple supervised classifiers were trained as baselines and candidates for "
        "deployment: Random Forest, Gradient Boosting, Linear Support Vector Machine (SVM), "
        "Logistic Regression, Decision Tree, Multinomial Naive Bayes, and K-Nearest "
        "Neighbours (K-NN). An MLP neural network configuration was also attempted but did "
        "not complete successfully in the experimental notebook run. Although the literature "
        "review discussed transformer models such as BERT as future-capable approaches, the "
        "implemented FYP focused on classical NLP pipelines that are lightweight, "
        "interpretable, and deployable on limited university hardware.",
    )
    add_body(
        doc,
        "Random Forest was selected as the best-performing model based on Accuracy, "
        "Precision, Recall, and F1-score on the held-out test set. The selected forest used "
        "100 estimators. The trained model and vectorizer were serialised with Joblib "
        "(crime_model.pkl and vectorizer.pkl) for production use.",
    )

    add_heading(doc, "3.7.2 Model Evaluation", level=3)
    add_body(
        doc,
        "Classification performance was measured using Accuracy, Precision, Recall, and "
        "F1-score. Complementary visualisations include a confusion matrix and comparative "
        "bar charts. Hyperparameters were controlled through consistent random seeds and "
        "shared train/test splits so that differences reflect model behaviour rather than "
        "sampling noise.",
    )
    add_table(
        doc,
        ["Model", "Accuracy", "Precision", "Recall", "F1-Score"],
        [
            ["Random Forest", "0.8955", "0.8955", "0.8955", "0.8955"],
            ["Gradient Boosting", "0.8840", "0.8858", "0.8840", "0.8839"],
            ["Linear SVM", "0.8635", "0.8636", "0.8635", "0.8635"],
            ["Logistic Regression", "0.8580", "0.8583", "0.8580", "0.8580"],
            ["Decision Tree", "0.8475", "0.8475", "0.8475", "0.8475"],
            ["Naive Bayes", "0.8305", "0.8316", "0.8305", "0.8304"],
            ["K-NN", "0.7830", "0.8251", "0.7830", "0.7757"],
        ],
        caption="Table 3.3: Model Evaluation Results on the Test Set",
    )

    add_heading(doc, "3.7.3 Model Deployment Framework", level=3)
    add_body(
        doc,
        "The selected Random Forest model is served through a Flask REST API that exposes "
        "health checks and a /predict endpoint. The Express backend proxies classification "
        "requests to the Flask service, enriches responses with blacklist matches, and "
        "stores outcomes in MongoDB. This separation allows the NLP model to be updated "
        "independently of the web interface.",
    )
    add_figure(
        doc,
        FIG / "fig3_2_pipeline.png",
        "Figure 3.3: NLP Crime Classification Pipeline",
        width=4.7,
    )

    add_heading(doc, "3.8 Requirements", level=2)
    add_heading(doc, "3.8.1 System Requirements", level=3)
    add_heading(doc, "Hardware Requirements", level=3)
    add_body(
        doc,
        "The hardware requirements define the minimum physical resources needed to develop, "
        "train, and run the application efficiently.",
    )
    add_table(
        doc,
        ["Component", "Minimum Specification"],
        [
            ["Processor", "Intel Core i5 (or equivalent) and above"],
            ["Memory (RAM)", "8 GB minimum; 16 GB recommended for training"],
            ["Storage", "20 GB free disk space"],
            ["Display", "1366 × 768 or higher"],
            ["Network", "Stable internet connection for APIs and scraping"],
        ],
        caption="Table 3.4: Hardware Requirements",
    )

    add_heading(doc, "Software Requirements", level=3)
    add_table(
        doc,
        ["Category", "Tools / Technologies"],
        [
            ["Operating Systems", "Windows 10/11, Linux, or macOS"],
            ["IDE / Editors", "Visual Studio Code, Jupyter Notebook"],
            ["Programming Languages", "JavaScript (Node.js), Python 3.x"],
            ["Frontend", "React 19, Vite, Tailwind CSS, Axios, Recharts"],
            ["Backend", "Express 5, Mongoose, JWT, Multer, Puppeteer, Cheerio"],
            ["ML / NLP", "scikit-learn, pandas, numpy, joblib, Flask"],
            ["Database", "MongoDB"],
            ["Version Control", "Git"],
        ],
        caption="Table 3.5: Software Requirements",
    )

    add_heading(doc, "3.8.2 User Requirements", level=3)
    add_heading(doc, "3.8.2.1 Functional Requirements", level=3)
    add_bullet(doc, "FR1: The system shall allow users to submit text for crime classification.")
    add_bullet(doc, "FR2: The system shall accept URL and file inputs and extract readable text.")
    add_bullet(doc, "FR3: The system shall return prediction label and confidence score.")
    add_bullet(doc, "FR4: The system shall authenticate users and enforce role-based access.")
    add_bullet(doc, "FR5: The system shall create investigation cases for crime detections.")
    add_bullet(doc, "FR6: The system shall allow investigators to claim and update cases.")
    add_bullet(doc, "FR7: The system shall support blacklist CRUD and monitoring scans.")
    add_bullet(doc, "FR8: The system shall notify relevant users when crimes are detected.")
    add_bullet(doc, "FR9: The system shall provide an admin dashboard and audit logs.")
    add_bullet(doc, "FR10: The system shall allow flagging of false or malicious reports.")

    add_heading(doc, "3.8.2.2 Non-Functional Requirements", level=3)
    add_bullet(doc, "NFR1 (Performance): Classification responses should return within a few seconds under normal load.")
    add_bullet(doc, "NFR2 (Usability): Interfaces should be clear for non-technical investigators.")
    add_bullet(doc, "NFR3 (Reliability): Analysis history and case records must be persisted durably.")
    add_bullet(doc, "NFR4 (Security): Passwords are hashed; API routes are JWT-protected where required.")
    add_bullet(doc, "NFR5 (Scalability): Services are modular so model and API can scale separately.")
    add_bullet(doc, "NFR6 (Maintainability): Code is organised by controllers, models, routes, and services.")
    add_bullet(doc, "NFR7 (Portability): The stack runs on common desktop operating systems.")

    add_heading(doc, "3.9 Feasibility Study", level=2)
    add_heading(doc, "3.9.1 Technical Feasibility", level=3)
    add_body(
        doc,
        "The project is technically feasible because all required tools are open-source and "
        "widely documented. React, Express, Flask, scikit-learn, and MongoDB are mature "
        "technologies already familiar to Computer Applications students. Model training on "
        "approximately 10,000 text samples is computationally practical on a standard laptop.",
    )
    add_heading(doc, "3.9.2 Economic Feasibility", level=3)
    add_body(
        doc,
        "Development relied on free software stacks and publicly available dataset construction "
        "effort. No commercial NLP API subscription is required for the core classifier, which "
        "makes the solution economically suitable for low-resource institutional settings.",
    )
    add_heading(doc, "3.9.3 Operational Feasibility", level=3)
    add_body(
        doc,
        "Operationally, BAREAI maps directly onto existing investigative workflows: intake, "
        "triage, assignment, and reporting. Role-based screens reduce training overhead for "
        "admins and investigators.",
    )
    add_heading(doc, "3.9.4 Schedule Feasibility", level=3)
    add_body(
        doc,
        "The work was organised into literature review, data preparation, model experimentation, "
        "backend/frontend development, integration testing, and thesis writing—aligned with the "
        "standard FYP academic calendar.",
    )

    add_heading(doc, "3.10 System Design", level=2)
    add_body(
        doc,
        "System design artefacts include a use-case view of actors and major interactions, and "
        "a logical database schema describing MongoDB collections and relationships. Detailed "
        "implementation screenshots and runtime results are presented in Chapter IV.",
    )
    add_heading(doc, "3.10.1 Use Case Diagram", level=3)
    add_body(
        doc,
        "The primary actors are Admin, Investigator, Public User, and the AI Service. Major "
        "use cases include classify text report, manage blacklist, claim/investigate case, "
        "view dashboard, manage users, flag false reports, and public analysis.",
    )
    add_figure(
        doc,
        FIG / "fig3_3_usecase.png",
        "Figure 3.4: Use Case Diagram of BAREAI",
        width=4.7,
    )

    add_heading(doc, "3.10.2 Database Design", level=3)
    add_body(
        doc,
        "MongoDB collections were designed around operational entities. User stores identity "
        "and role fields. History stores each analysis event. InvestigationCase links one-to-one "
        "with a crime History record and tracks assignment status. BlacklistItem defines "
        "monitored sources. Notification and ActivityLog support communication and auditing. "
        "InvestigationReport stores investigator findings. Relationships are enforced logically "
        "through foreign-key style ObjectId references.",
    )
    add_figure(
        doc,
        FIG / "fig3_4_database.png",
        "Figure 3.5: Logical Database Schema of BAREAI",
        width=4.7,
    )
    add_table(
        doc,
        ["Collection", "Key Fields", "Purpose"],
        [
            ["User", "email, role, account_status", "Authentication and RBAC"],
            ["History", "content, prediction, confidence", "Analysis archive"],
            ["InvestigationCase", "history, status, assignedOfficer", "Case workflow"],
            ["InvestigationReport", "case, findings, status", "Formal reports"],
            ["BlacklistItem", "type, value, priority", "Watchlist monitoring"],
            ["Notification", "recipient, type, read", "In-app alerts"],
            ["ActivityLog", "action, module, user", "Audit trail"],
        ],
        caption="Table 3.6: Database Collections Summary",
    )
    inject_expansions(doc, "ch3_extra", add_body)

    if through_chapter <= 3:
        return

    # ==================================================================
    # CHAPTER IV
    # ==================================================================
    add_page_break(doc)
    add_heading(doc, "CHAPTER FOUR: IMPLEMENTATION AND RESULTS", level=1)

    add_heading(doc, "4.1 Dataset Description", level=2)
    add_body(
        doc,
        "The dataset used for training and evaluation consists of 9,999 Somali-oriented text "
        "entries organised as rows with URL, text, and category attributes. Labels are binary "
        "and nearly balanced, which reduces the need for aggressive resampling. Examples "
        "include news-style narratives and social posts discussing incidents that may or may "
        "not constitute criminal activity. Average document length is about 277 words, with "
        "substantial variance, reflecting real-world textual heterogeneity.",
    )
    add_figure(
        doc,
        FIG / "fig4_3_class_distribution.png",
        "Figure 4.1: Dataset Class Distribution (N = 9,999)",
        width=4.5,
    )

    add_heading(doc, "4.2 ML Pipeline Implementation", level=2)
    add_body(
        doc,
        "The machine learning pipeline was developed through a structured process that included "
        "data loading, cleaning, stopword removal, TF-IDF vectorisation, stratified splitting, "
        "multi-model training, metric comparison, model selection, serialisation, and API "
        "deployment. Training and testing were performed on disjoint subsets to assess "
        "generalisation. The final Random Forest classifier was validated on the 2,000-sample "
        "test set and then packaged for Flask inference.",
    )
    add_body(
        doc,
        "At runtime, the Flask service transforms incoming text with the saved TF-IDF "
        "vectorizer, obtains class probabilities from Random Forest, and applies Somali "
        "crime-keyword overrides when explicit high-risk terms are present. Location mentions "
        "are extracted with pattern matching against known Somali place names. The Express "
        "layer stores the enriched response and triggers case creation when isCrime is true.",
    )
    add_figure(
        doc,
        FIG / "fig4_4_dataflow.png",
        "Figure 4.2: End-to-End Classification and Case Dispatch Flow",
        width=4.7,
    )
    add_figure(
        doc,
        FIG / "fig3_2_pipeline.png",
        "Figure 4.3: ML Pipeline Implementation Stages",
        width=4.7,
    )

    add_heading(doc, "4.3 System Design and Deployment", level=2)
    add_heading(doc, "4.3.1 Backend Implementation", level=3)
    add_body(
        doc,
        "The backend is implemented in Node.js using Express. Key route groups include "
        "/api/auth, /api/analysis, /api/investigation, /api/blacklist, /api/dashboard, "
        "/api/notifications, /api/reports, /api/users, and /api/audit-logs. Controllers "
        "coordinate validation, AI inference calls, MongoDB persistence, and notification "
        "dispatch. Background Facebook monitoring uses Puppeteer; website scanning uses "
        "Cheerio with SSRF-safe fetching.",
    )

    add_heading(doc, "4.3.2 Frontend Implementation", level=3)
    add_body(
        doc,
        "The frontend is a React SPA with role-aware routing. Public users access landing and "
        "analysis pages. Investigators focus on cases, blacklist tools, and reports. Admins "
        "access the dashboard, user management, and audit logs. Charts are rendered with "
        "Recharts, and PDF exports are supported where reporting is required.",
    )

    add_heading(doc, "4.3.3 AI Service Integration", level=3)
    add_body(
        doc,
        "The Flask service listens on port 5001 by default and exposes /health and /predict. "
        "The Express backend communicates with AI_MODEL_URL, keeping the browser isolated "
        "from direct model access. This improves security and allows independent redeployment "
        "of the classifier.",
    )

    add_heading(doc, "4.3.4 Deployment Topology", level=3)
    add_body(
        doc,
        "In the local development and demonstration environment, the frontend runs on port "
        "5173, the Express API on port 5000, and the Flask model service on port 5001, with "
        "MongoDB as the shared database. The same topology can be migrated to a cloud host "
        "by containerising each service.",
    )
    add_figure(
        doc,
        FIG / "fig4_5_deployment.png",
        "Figure 4.4: Local Deployment Topology of BAREAI Services",
        width=4.7,
    )

    add_heading(doc, "4.3.5 Implemented User Interface Modules", level=3)
    add_body(
        doc,
        "The following interface modules were implemented and tested during system integration:",
    )
    add_table(
        doc,
        ["Module", "Primary Users", "Description"],
        [
            ["Landing / Public Analysis", "Public users", "Submit text/URL/file for classification"],
            ["Admin Dashboard", "Admin", "Operational KPIs and trends"],
            ["User Management", "Admin", "Create and manage accounts/roles"],
            ["Case Management", "Admin, Investigator", "Claim, update, and close cases"],
            ["Blacklist / Monitoring", "Admin, Investigator", "Watch Facebook/websites/keywords"],
            ["Reports & Audit Logs", "Admin, Investigator", "Accountability and analytics"],
            ["Profile & Settings", "All authenticated", "Account and preference management"],
        ],
        caption="Table 4.1: Implemented Interface Modules",
    )
    add_body(
        doc,
        "Representative interface screens of the implemented website are presented below. "
        "Each screen demonstrates how users interact with Text, File, URL, and Batch analysis "
        "features and how outputs are shown as crime-related or non-crime-related.",
    )
    add_figure(doc, FIG / "fig4_10_landing.png", "Figure 4.5: BAREAI Landing / Home Page", width=4.7)
    add_figure(doc, FIG / "fig4_11_login.png", "Figure 4.6: BAREAI Login Page", width=4.7)
    add_figure(doc, FIG / "fig4_8_analysis_ui.png", "Figure 4.7: Text Analysis Page with Crime / Non-Crime Result", width=4.7)
    add_figure(doc, FIG / "fig4_13_url_analysis.png", "Figure 4.8: URL Analysis Page", width=4.7)
    add_figure(doc, FIG / "fig4_14_file_analysis.png", "Figure 4.9: File Analysis Page", width=4.7)
    add_figure(doc, FIG / "fig4_15_batch_analysis.png", "Figure 4.10: Batch Analysis Page", width=4.7)
    add_figure(doc, FIG / "fig4_7_dashboard_ui.png", "Figure 4.11: Administrator Dashboard", width=4.7)
    add_figure(doc, FIG / "fig4_9_cases_ui.png", "Figure 4.12: Investigation Case Management Page", width=4.7)
    add_figure(doc, FIG / "fig4_12_blacklist.png", "Figure 4.13: Blacklist and Monitoring Page", width=4.7)
    add_figure(doc, FIG / "fig4_16_notifications_reports.png", "Figure 4.14: Notifications and Reports Page", width=4.7)

    add_heading(doc, "4.4 Website Interface Screenshots — Feature Walkthrough", level=2)
    add_body(
        doc,
        "The landing page introduces BAREAI as an intelligence platform for automatic "
        "classification of social media posts. It states clearly that the system separates "
        "content into crime-related and non-crime-related classes and highlights Text, File, "
        "URL, and Batch analysis entry points. The login page authenticates Admin, "
        "Investigator, and Public User roles before granting access to protected modules.",
    )
    add_body(
        doc,
        "On the Text Analysis page, a user pastes a social media post; the system returns "
        "the predicted class and confidence. URL Analysis fetches webpage text before "
        "classification. File Analysis accepts PDF/DOCX/TXT uploads. Batch Analysis lists "
        "several posts and labels each one independently. These four features together "
        "cover the most common ways crime-related narratives appear online.",
    )
    add_body(
        doc,
        "Administrative screens support operations after classification. The dashboard "
        "summarises analyses and detections; the case board allows investigators to claim "
        "crime-related items; blacklist monitoring watches high-risk Facebook pages and "
        "websites; and notifications/reports keep the organisation informed. Together, "
        "these screens show that BAREAI is not only a model demo but a complete workflow "
        "application.",
    )

    add_heading(doc, "4.5 Data Visualization (Results Presentation)", level=2)
    add_body(
        doc,
        "Experimental results are summarised using comparative charts, a confusion matrix, "
        "and metric bars for the best model. Random Forest achieved 89.55% across Accuracy, "
        "Precision, Recall, and F1-score on the test set. Approximate confusion-matrix cell "
        "counts for a balanced 2,000-sample test partition are shown to illustrate error "
        "distribution under near-symmetric performance.",
    )
    add_figure(
        doc,
        FIG / "fig4_1_model_comparison.png",
        "Figure 4.15: Comparative Performance of Classification Models",
        width=4.7,
    )
    add_figure(
        doc,
        FIG / "fig4_2_confusion_matrix.png",
        "Figure 4.16: Confusion Matrix of the Random Forest Classifier",
        width=4.6,
    )
    add_figure(
        doc,
        FIG / "fig4_6_rf_metrics.png",
        "Figure 4.17: Evaluation Metrics of the Selected Random Forest Model",
        width=4.7,
    )

    add_heading(doc, "4.6 Model Comparison", level=2)
    add_body(
        doc,
        "Table 4.2 consolidates the comparative evaluation. Random Forest outperformed all "
        "baselines, exceeding Gradient Boosting by about 1.15 percentage points in accuracy "
        "and outperforming Linear SVM by about 3.20 percentage points. Tree ensembles were "
        "particularly effective for TF-IDF features of Somali crime text, while K-NN lagged "
        "due to high-dimensional sparse representations.",
    )
    add_table(
        doc,
        ["Rank", "Model", "Accuracy (%)", "F1-Score (%)", "Remark"],
        [
            ["1", "Random Forest", "89.55", "89.55", "Selected for deployment"],
            ["2", "Gradient Boosting", "88.40", "88.39", "Strong runner-up"],
            ["3", "Linear SVM", "86.35", "86.35", "Competitive linear baseline"],
            ["4", "Logistic Regression", "85.80", "85.80", "Fast interpretable baseline"],
            ["5", "Decision Tree", "84.75", "84.75", "Weaker generalisation"],
            ["6", "Naive Bayes", "83.05", "83.04", "Strong speed, lower score"],
            ["7", "K-NN", "78.30", "77.57", "Least suitable"],
        ],
        caption="Table 4.2: Ranked Model Comparison Summary",
    )
    add_body(
        doc,
        "Beyond offline metrics, system-level testing confirmed that the deployed hybrid "
        "pipeline (Random Forest + TF-IDF + keyword enrichment) successfully classifies live "
        "inputs and integrates with case management. This demonstrates that the research "
        "objectives—reducing manual triage, preprocessing crime text, extracting features, "
        "and delivering an automated NLP system—were achieved in a working prototype.",
    )
    inject_expansions(doc, "ch4_extra", add_body)

    add_heading(doc, "4.7 Detailed Testing Scenarios", level=2)
    add_body(
        doc,
        "To strengthen validation evidence, structured testing scenarios were executed across "
        "the four analysis features. Each scenario records the input type, expected class, "
        "and observed system behaviour. The scenarios confirm that BAREAI consistently "
        "returns crime-related or non-crime-related labels with confidence scores.",
    )
    add_table(
        doc,
        ["Scenario", "Input Feature", "Example Content", "Expected Class", "Observed Behaviour"],
        [
            ["S1", "Text Analysis", "Post describing an armed attack", "Crime-related", "Correct label + case created"],
            ["S2", "Text Analysis", "Sports match commentary", "Non-crime-related", "Correct label, no case"],
            ["S3", "URL Analysis", "News page on explosion", "Crime-related", "Text extracted and classified"],
            ["S4", "File Analysis", "PDF incident report", "Crime-related", "Text extracted and classified"],
            ["S5", "Batch Analysis", "Mixed set of 5 posts", "Mixed", "Per-item labels returned"],
            ["S6", "Blacklist Monitor", "Watched Facebook page post", "Crime-related", "Alert + notification"],
            ["S7", "Case Claim", "Pending crime case", "N/A", "First investigator claims successfully"],
            ["S8", "Auth Guard", "Unauthenticated admin URL", "N/A", "Access blocked by JWT"],
        ],
        caption="Table 4.3: Detailed Functional Testing Scenarios",
    )
    add_body(
        doc,
        "Scenario S1–S5 validate the core NLP path. Scenario S6 validates monitoring "
        "integration. Scenario S7 validates investigation workflow. Scenario S8 validates "
        "security controls. Together they show that classification quality and application "
        "reliability were both examined before thesis submission.",
    )
    add_body(
        doc,
        "Error-handling tests included empty text submission, unsupported file types, and "
        "unreachable model service. In each case the API returned controlled error messages "
        "rather than crashing the frontend. These defensive behaviours are important for "
        "demonstration stability and future production hardening.",
    )
    add_body(
        doc,
        "Usability observations during team walkthroughs indicated that investigators "
        "understood the crime-related versus non-crime-related labels quickly, while admins "
        "valued dashboard aggregates. Minor suggestions included larger confidence "
        "visualisation and exportable CSV history; these are noted under future work.",
    )

    # ==================================================================
    # CHAPTER V
    # ==================================================================
    add_page_break(doc)
    add_heading(doc, "CHAPTER FIVE: DISCUSSION AND CONCLUSION", level=1)

    add_heading(doc, "5.1 Introduction", level=2)
    add_body(
        doc,
        "This chapter interprets the experimental and implementation findings of the study, "
        "relates them to the research objectives and questions, discusses contributions, "
        "acknowledges limitations, and presents recommendations for future work. It concludes "
        "with the main takeaway of the BAREAI project.",
    )

    add_heading(doc, "5.2 Interpretation of Results", level=2)
    add_body(
        doc,
        "The study set out to develop an intelligent automated framework that uses machine "
        "learning and NLP to classify crime-related text reports. The best model, Random "
        "Forest, reached 89.55% Accuracy/Precision/Recall/F1 on a balanced held-out test set. "
        "This indicates that classical TF-IDF features, when paired with an ensemble "
        "classifier, can capture discriminative lexical patterns in Somali crime-related "
        "narratives without requiring large transformer models.",
    )
    add_body(
        doc,
        "In relation to the research questions: (1) manual classification reliance is reduced "
        "by providing an automated prediction with confidence and case dispatch; (2) data were "
        "successfully collected and preprocessed through cleaning, stopword removal, and "
        "normalisation; (3) TF-IDF feature extraction produced representations that support "
        "accurate classification; and (4) a complete NLP-enabled web system was developed and "
        "integrated with monitoring and investigation workflows.",
    )
    add_body(
        doc,
        "Compared with weaker baselines such as K-NN (78.30% accuracy), the selected model "
        "offers a practically meaningful improvement for triage. The near equality of "
        "precision and recall suggests balanced behaviour across crime and non-crime classes, "
        "which is desirable when both false alarms and missed crimes are costly.",
    )

    add_heading(doc, "5.3 Contribution to Knowledge", level=2)
    add_body(
        doc,
        "The main contributions of this FYP are as follows:",
    )
    add_bullet(
        doc,
        "A Somali-oriented labelled corpus of nearly 10,000 texts for binary crime-related "
        "classification.",
    )
    add_bullet(
        doc,
        "An empirical comparison of seven classical classifiers showing Random Forest as the "
        "strongest performer at 89.55% F1-score.",
    )
    add_bullet(
        doc,
        "A production-style three-tier architecture (React–Express–Flask–MongoDB) that connects "
        "NLP inference to operational investigation workflows.",
    )
    add_bullet(
        doc,
        "Practical monitoring features (Facebook/website blacklist scanning) that extend "
        "classification beyond one-off manual paste inputs.",
    )
    add_bullet(
        doc,
        "A case-claiming and notification mechanism that demonstrates how AI predictions can "
        "support law-enforcement triage in low-resource settings.",
    )

    add_heading(doc, "5.4 Challenges and Limitations", level=2)
    add_body(
        doc,
        "Several challenges and limitations were encountered. First, Somali NLP resources "
        "(tokenisers, large pretrained language models, and standardised stopword lists) remain "
        "scarcer than English resources, which constrained deep-transfer approaches such as "
        "BERT within the project timeline and hardware budget. Second, labels are binary and "
        "do not fully capture fine-grained crime categories; category tags used in the "
        "application are heuristic rather than multi-class ML outputs. Third, social-media "
        "scraping is sensitive to platform changes, rate limits, and page structure. Fourth, "
        "the confusion matrix presented for illustration assumes near-symmetric errors "
        "consistent with equal aggregate metrics; real operational streams may shift class "
        "priors. Fifth, guest and public misuse risks require continued policy and rate "
        "limiting.",
    )

    add_heading(doc, "5.5 Future Work", level=2)
    add_body(
        doc,
        "Future research and development may extend BAREAI in the following directions:",
    )
    add_bullet(doc, "Increase dataset size and diversity (dialects, code-mixing, informal chat text).")
    add_bullet(doc, "Train multi-class crime-type models (homicide, theft, cybercrime, etc.).")
    add_bullet(doc, "Evaluate transformer-based Somali/multilingual encoders (e.g., BERT variants) when compute allows.")
    add_bullet(doc, "Improve explainability with feature-importance or token-level highlights for investigators.")
    add_bullet(doc, "Deploy on secure cloud infrastructure with continuous model monitoring and retraining.")
    add_bullet(doc, "Integrate with official policing databases under appropriate legal and ethical frameworks.")
    add_bullet(doc, "Strengthen adversarial robustness against obfuscated crime language.")

    add_heading(doc, "5.6 Final Conclusion", level=2)
    add_body(
        doc,
        "This study designed, implemented, and evaluated BAREAI, an AI-based system for "
        "automatic classification of crime-related text reports using NLP. Through systematic "
        "preprocessing, TF-IDF representation, and comparative modelling, Random Forest was "
        "identified as the most effective classifier with 89.55% test performance. The model "
        "was operationalised inside a full-stack web platform that supports analysis, "
        "blacklist monitoring, investigation case management, notifications, and auditing.",
    )
    add_body(
        doc,
        "The take-home contribution is practical: Somali crime-text triage can be accelerated "
        "by a lightweight, transparent ML pipeline embedded in investigator-facing software. "
        "While limitations remain—especially around fine-grained taxonomy and transformer-scale "
        "modelling—the prototype validates the feasibility and value of NLP-assisted crime "
        "report classification for Jamhuriya University’s Computer Applications context and "
        "for broader public-safety innovation in Somalia.",
    )
    add_body(
        doc,
        "In summary, the research objectives were met: reliance on purely manual sorting was "
        "reduced; crime-related text data were collected and preprocessed; discriminative "
        "features were extracted and analysed; and an automated NLP system was developed, "
        "tested, and demonstrated as a working FYP product.",
    )
    inject_expansions(doc, "ch5_extra", add_body)
    add_body(
        doc,
        "As a closing synthesis, the project affirms three practical claims. First, social media "
        "posts can be analysed at scale when NLP pipelines are paired with clear binary targets. "
        "Second, Random Forest with TF-IDF can deliver near-90% test accuracy for Somali-oriented "
        "crime text under the dataset conditions studied here. Third, a web application that "
        "exposes Text, File, URL, and Batch analysis makes those research results usable for "
        "investigators rather than leaving them in notebooks alone.",
    )
    add_body(
        doc,
        "Therefore, the thesis recommends pilot adoption inside academic and institutional labs "
        "that can supervise ethical use, followed by incremental upgrades in data volume, "
        "multilingual coverage, and explainability. With those steps, BAREAI can evolve from an "
        "FYP prototype into a sustained early-warning assistant for digital crime triage.",
    )

    # Appendix figures index
    add_page_break(doc)
    add_heading(doc, "APPENDIX A: FIGURE INDEX", level=1)
    add_bullet(doc, "Figure 3.0: Multi-Source Input Analysis Features")
    add_bullet(doc, "Figure 3.1: BAREAI System Architecture Pipeline")
    add_bullet(doc, "Figure 3.2: Text Preprocessing Flowchart")
    add_bullet(doc, "Figure 3.3: NLP Crime Classification Pipeline")
    add_bullet(doc, "Figure 3.4: Use Case Diagram of BAREAI")
    add_bullet(doc, "Figure 3.5: Logical Database Schema of BAREAI")
    add_bullet(doc, "Figure 4.1: Dataset Class Distribution (N = 9,999)")
    add_bullet(doc, "Figure 4.2: End-to-End Classification and Case Dispatch Flow")
    add_bullet(doc, "Figure 4.3: ML Pipeline Implementation Stages")
    add_bullet(doc, "Figure 4.4: Local Deployment Topology of BAREAI Services")
    add_bullet(doc, "Figure 4.5: BAREAI Landing / Home Page")
    add_bullet(doc, "Figure 4.6: BAREAI Login Page")
    add_bullet(doc, "Figure 4.7: Text Analysis Page with Crime / Non-Crime Result")
    add_bullet(doc, "Figure 4.8: URL Analysis Page")
    add_bullet(doc, "Figure 4.9: File Analysis Page")
    add_bullet(doc, "Figure 4.10: Batch Analysis Page")
    add_bullet(doc, "Figure 4.11: Administrator Dashboard")
    add_bullet(doc, "Figure 4.12: Investigation Case Management Page")
    add_bullet(doc, "Figure 4.13: Blacklist and Monitoring Page")
    add_bullet(doc, "Figure 4.14: Notifications and Reports Page")
    add_bullet(doc, "Figure 4.15: Comparative Performance of Classification Models")
    add_bullet(doc, "Figure 4.16: Confusion Matrix of the Random Forest Classifier")
    add_bullet(doc, "Figure 4.17: Evaluation Metrics of the Selected Random Forest Model")

    add_heading(doc, "APPENDIX B: TABLE INDEX", level=1)
    add_bullet(doc, "Table 3.1: Dataset Column Description and Summary Statistics")
    add_bullet(doc, "Table 3.2: Feature Engineering and Split Configuration")
    add_bullet(doc, "Table 3.3: Model Evaluation Results on the Test Set")
    add_bullet(doc, "Table 3.4: Hardware Requirements")
    add_bullet(doc, "Table 3.5: Software Requirements")
    add_bullet(doc, "Table 3.6: Database Collections Summary")
    add_bullet(doc, "Table 4.1: Implemented Interface Modules")
    add_bullet(doc, "Table 4.2: Ranked Model Comparison Summary")
    add_bullet(doc, "Table 4.3: Detailed Functional Testing Scenarios")

    add_page_break(doc)
    add_heading(doc, "APPENDIX C: SAMPLE CLASSIFICATION EXAMPLES", level=1)
    add_body(
        doc,
        "This appendix lists illustrative social-media-style inputs and the expected binary "
        "class used during manual review and system demonstration. Examples are shortened "
        "for readability and are intended to show how crime-related and non-crime-related "
        "posts differ in language.",
    )
    samples = [
        ["E01", "weerar hubaysan oo ka dhacay degmada X", "Crime-related"],
        ["E02", "qarax ka dhacay suuqa Y oo dhaawacay dad", "Crime-related"],
        ["E03", "tuugo baabuur habeenkii oo la xaday", "Crime-related"],
        ["E04", "afduub lagu haysto shaqaale dawladeed", "Crime-related"],
        ["E05", "dil lagu geystay xaafadda Z", "Crime-related"],
        ["E06", "kufsi iyo xadgudub lagu soo sheegay", "Crime-related"],
        ["E07", "baasaboorka been abuurka ah ayaa la qabtay", "Crime-related"],
        ["E08", "koox tuugo ah oo weerartay dukaan", "Crime-related"],
        ["E09", "boolisku waxay xireen tuhmanayaal", "Crime-related"],
        ["E10", "qalalaase hubaysan oo ka taagan wadada", "Crime-related"],
        ["E11", "ciyaar fudud oo ka dhacday garoonka", "Non-crime-related"],
        ["E12", "hecdo cusub oo fannaanka soo saaray", "Non-crime-related"],
        ["E13", "roobab mahiigaan ah oo beeraha ka caawiyay", "Non-crime-related"],
        ["E14", "shirkad cusub oo shaqo furatay", "Non-crime-related"],
        ["E15", "ardayda jaamacadda waxay galeen imtixaan", "Non-crime-related"],
        ["E16", "bandhig farshaxaneed oo dadku daawadeen", "Non-crime-related"],
        ["E17", "tartanka kubadda cagta ayaa dhammaaday", "Non-crime-related"],
        ["E18", "suuqa maanta waxaa ka jira qiime dhimis", "Non-crime-related"],
        ["E19", "buug cusub oo la daabacay ayaa soo baxay", "Non-crime-related"],
        ["E20", "kulanka beesha ayaa nabad ku dhammaaday", "Non-crime-related"],
        ["E21", "qorshe dawladeed oo waxbarasho ah", "Non-crime-related"],
        ["E22", "warbixin caafimaad oo tallaalka ku saabsan", "Non-crime-related"],
        ["E23", "dab qabtay guri oo dad ku dhinteen", "Crime-related"],
        ["E24", "baabuur isku dhacey oo dhaawac keenay", "Crime-related"],
        ["E25", "hanjabaad ah in la weeraro goob dadwayne", "Crime-related"],
    ]
    add_table(
        doc,
        ["ID", "Sample Text (shortened)", "Class"],
        samples,
        caption="Table C.1: Sample Crime-related and Non-crime-related Texts",
    )
    add_body(
        doc,
        "These samples are not a substitute for the full training corpus; they illustrate the "
        "semantic contrast the classifier must learn. During demonstration, similar texts were "
        "submitted through Text Analysis and Batch Analysis screens to verify that predicted "
        "labels matched reviewer expectations in most cases.",
    )
    add_body(
        doc,
        "Additional demonstration notes: when a crime-related sample was submitted, investigators "
        "received a case notification and could claim the item. When a non-crime-related sample "
        "was submitted, history was stored without creating an investigation case. This behaviour "
        "matches the operational rule that only crime-related detections enter the case queue.",
    )
    add_body(
        doc,
        "Appendix C therefore complements Chapter Four by providing concrete textual examples "
        "that examiners can reuse during viva questioning. Students can paste any listed example "
        "into the Text Analysis page to reproduce the crime-related versus non-crime-related "
        "decision live.",
    )
    add_page_break(doc)
    add_heading(doc, "APPENDIX D: RECOMMENDED VIVA DEMONSTRATION SCRIPT", level=1)
    add_body(
        doc,
        "1. Open the BAREAI landing page and explain that the system classifies social media "
        "posts as crime-related or non-crime-related.",
    )
    add_body(
        doc,
        "2. Demonstrate Text Analysis with one crime-related and one non-crime-related sample "
        "from Appendix C.",
    )
    add_body(
        doc,
        "3. Demonstrate URL Analysis and File Analysis briefly, then Batch Analysis with mixed posts.",
    )
    add_body(
        doc,
        "4. Login as investigator, show a pending case, claim it, and update status.",
    )
    add_body(
        doc,
        "5. Login as admin, show dashboard totals, blacklist monitoring, and audit logs.",
    )
    add_body(
        doc,
        "6. Conclude by stating that Random Forest (~89.55% accuracy) powers the Flask NLP "
        "service behind the website features.",
    )
    add_body(
        doc,
        "The demonstration script above is intentionally short so that the viva can remain "
        "within the allocated time while still covering classification quality, website "
        "features, and operational workflow. Teams should rehearse once with the backend, "
        "frontend, and Flask model services running together.",
    )
    add_body(
        doc,
        "If network or browser issues occur during viva, Appendix C samples and Chapter Four "
        "screenshots remain available as fallback evidence that the system was implemented and "
        "tested. Examiners can still inspect the architecture pipeline and model comparison "
        "tables in the main chapters.",
    )
    add_body(
        doc,
        "This closes the supplementary material before the reference list. The following "
        "extended discussion notes provide additional reflective depth for examination "
        "readiness while preserving the main chapter narrative.",
    )

    add_page_break(doc)
    add_heading(doc, "APPENDIX E: EXTENDED DISCUSSION NOTES", level=1)
    for note in [
        "This appendix consolidates extended discussion notes that support Chapters Four and Five. It restates that BAREAI analyses social media posts and related digital reports, then separates them into crime-related and non-crime-related classes.",
        "The four analysis features—Text Analysis, File Analysis, URL Analysis, and Batch Analysis—were implemented to mirror real investigator intake channels. Each feature ultimately calls the same Random Forest NLP service so that class semantics remain consistent.",
        "Accuracy near 89.55 percent should be interpreted in context: the test set was balanced and stratified. In live streams where non-crime posts dominate, precision-oriented thresholds or calibrated probabilities may be preferred.",
        "Operationally, case claiming prevents duplicated investigator effort. Notifications reduce missed detections. Audit logs support accountability. These application features turn a classifier into a usable security assistant.",
        "Ethically, automated labels must remain advisory. Human review is required before coercive action. The thesis therefore positions BAREAI as decision support, not autonomous judgement.",
        "Technically, isolating Flask inference enables independent model upgrades. Future teams can replace Random Forest with stronger encoders without rewriting React screens, provided input/output contracts remain stable.",
        "From a JUST assessment perspective, the project demonstrates problem framing, literature awareness, methodology, working software, evaluation tables, and reflective conclusions—the complete AI FYP package.",
        "Students maintaining the repository should keep dataset provenance, model training notebooks, and environment files under version control so that viva claims remain reproducible.",
        "In closing, Appendix E exists to provide additional reflective depth requested for thesis length and examination readiness while preserving the clarity of Chapters One to Five.",
        "Readers seeking implementation detail should return to Chapter Three for methodology and Chapter Four for screenshots, metrics, and testing scenarios.",
    ]:
        add_body(doc, note)

    add_page_break(doc)
    add_heading(doc, "REFERENCES", level=1)
    data = json.loads(CH12_JSON.read_text(encoding="utf-8"))
    ref_paras = data["sections"].get("REFERENCES", [])
    for para in ref_paras:
        chunks = re.split(r"(?<=\.)\s+(?=[A-Z][a-z]+,\s+[A-Z])", para)
        for ch in chunks:
            ch = ch.strip()
            if len(ch) > 40:
                add_paragraph(
                    doc,
                    ch,
                    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    line_spacing=1.0,
                    space_after=8,
                    first_line_indent=0,
                )
    add_paragraph(
        doc,
        "Faculty of Computer and Information Technology, Jamhuriya University of Science and "
        "Technology. (n.d.). Undergraduate researcher’s compass: A practical guide for final "
        "year computer science students (AI based FYPs).",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        line_spacing=1.0,
        space_after=8,
        first_line_indent=0,
    )


def add_ch345_cover_and_contents(doc):
    """Standalone cover + CONTENTS for Chapters III–V only."""
    for _ in range(2):
        add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    add_paragraph(
        doc,
        PROJECT_TITLE,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.5,
        space_after=8,
    )
    add_paragraph(
        doc,
        "CHAPTERS III, IV & V",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.5,
        space_after=4,
    )
    add_paragraph(
        doc,
        "Methodology | Implementation and Results | Discussion and Conclusion",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.5,
        space_after=12,
    )
    for name, _cid in CANDIDATES:
        add_paragraph(doc, name, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, space_after=2)
    add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    add_paragraph(
        doc,
        "JAMHURIYA UNIVERSITY OF SCIENCE AND TECHNOLOGY (JUST)",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.5,
        space_after=2,
    )
    add_paragraph(
        doc,
        "FACULTY OF COMPUTER & INFORMATION TECHNOLOGY",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.5,
        space_after=2,
    )
    add_paragraph(doc, "AUGUST 2026", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    add_page_break(doc)

    add_centered_title(doc, "CONTENTS", outline_level=None)
    toc = [
        ("CHAPTER III: METHODOLOGY", "1", 0),
        ("3.1 Introduction", "1", 1),
        ("3.2 System Overview", "2", 1),
        ("3.3 System Features", "2", 1),
        ("3.4 System Architecture", "4", 1),
        ("3.5 Data Acquisition and Preprocessing", "5", 1),
        ("3.6 Feature Engineering / Representation", "6", 1),
        ("3.7 Model / Algorithm Development", "7", 1),
        ("3.8 Requirements", "9", 1),
        ("3.9 Feasibility Study", "11", 1),
        ("3.10 System Design", "12", 1),
        ("CHAPTER IV: IMPLEMENTATION AND RESULTS", "14", 0),
        ("4.1 Dataset Description", "14", 1),
        ("4.2 ML Pipeline Implementation", "15", 1),
        ("4.3 System Design and Deployment", "16", 1),
        ("4.4 Data Visualization (Results Presentation)", "19", 1),
        ("4.5 Model Comparison", "21", 1),
        ("CHAPTER V: DISCUSSION AND CONCLUSION", "23", 0),
        ("5.1 Introduction", "23", 1),
        ("5.2 Interpretation of Results", "23", 1),
        ("5.3 Contribution to Knowledge", "24", 1),
        ("5.4 Challenges and Limitations", "25", 1),
        ("5.5 Future Work", "25", 1),
        ("5.6 Final Conclusion", "26", 1),
        ("APPENDIX A: FIGURE INDEX", "28", 0),
        ("APPENDIX B: TABLE INDEX", "29", 0),
        ("REFERENCES", "30", 0),
    ]
    for title, page, level in toc:
        add_toc_entry(doc, title, page, level=level)
    add_page_break(doc)


def add_ch123_cover_and_contents(doc):
    """Standalone cover + CONTENTS for Chapters I–III only."""
    for _ in range(2):
        add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    add_paragraph(
        doc,
        PROJECT_TITLE,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.5,
        space_after=8,
    )
    add_paragraph(
        doc,
        "CHAPTERS I, II & III",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.5,
        space_after=4,
    )
    add_paragraph(
        doc,
        "Introduction | Literature Review | Methodology",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.5,
        space_after=12,
    )
    for name, _cid in CANDIDATES:
        add_paragraph(doc, name, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5, space_after=2)
    add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    add_paragraph(
        doc,
        "JAMHURIYA UNIVERSITY OF SCIENCE AND TECHNOLOGY (JUST)",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.5,
        space_after=2,
    )
    add_paragraph(
        doc,
        "FACULTY OF COMPUTER & INFORMATION TECHNOLOGY",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.5,
        space_after=2,
    )
    add_paragraph(doc, "AUGUST 2026", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    add_page_break(doc)

    add_centered_title(doc, "CONTENTS", outline_level=None)
    toc = [
        ("CHAPTER ONE: INTRODUCTION", "1", 0),
        ("1.1 Background of the Study", "1", 1),
        ("1.2 Problem Statement", "5", 1),
        ("1.3 Research Objectives", "7", 1),
        ("1.3.1 General Objective", "7", 2),
        ("1.3.2 Specific Objectives", "7", 2),
        ("1.4 Research Questions", "8", 1),
        ("1.5 Significance of the Study", "8", 1),
        ("1.6 Scope of the Study", "9", 1),
        ("1.7 Organization of the Study", "10", 1),
        ("CHAPTER TWO: LITERATURE REVIEW", "12", 0),
        ("2.1 Introduction", "12", 1),
        ("2.2 Impact of Crimes on Social Media", "16", 1),
        ("2.3 Traditional Classification Techniques in Crimes", "20", 1),
        ("2.4 Cybercrime Detection in Social Media", "24", 1),
        ("2.5 Challenges in the Classification of Crime-Related Text", "28", 1),
        ("2.6 Related Work", "32", 1),
        ("2.7 Proposed System", "36", 1),
        ("CHAPTER THREE: METHODOLOGY", "38", 0),
        ("3.1 Introduction", "38", 1),
        ("3.2 System Overview", "39", 1),
        ("3.3 System Features", "41", 1),
        ("3.4 System Architecture", "45", 1),
        ("3.5 Data Acquisition and Preprocessing", "48", 1),
        ("3.6 Feature Engineering / Representation", "51", 1),
        ("3.7 Model / Algorithm Development", "53", 1),
        ("3.8 Requirements", "57", 1),
        ("3.9 Feasibility Study", "60", 1),
        ("3.10 System Design", "62", 1),
    ]
    for title, page, level in toc:
        add_toc_entry(doc, title, page, level=level)
    add_page_break(doc)


def build_chapters_345_only():
    """Generate a separate document containing only Chapters III–V (connected)."""
    doc = Document()
    set_margins(doc.sections[0])
    configure_styles(doc)
    _set_section_page_number(doc.sections[0], start=1, roman=False, show=True)

    add_ch345_cover_and_contents(doc)
    add_chapters_three_to_five(doc, page_break_before=False, through_chapter=5)
    save_thesis(doc, mode="345")


def build_chapters_123_only():
    """Generate Chapters I–III only, with JUST formatting (justify + clean indent)."""
    doc = Document()
    set_margins(doc.sections[0])
    configure_styles(doc)
    _set_section_page_number(doc.sections[0], start=1, roman=False, show=True)

    add_ch123_cover_and_contents(doc)
    add_chapters_one_two(doc)
    add_chapters_three_to_five(doc, page_break_before=True, through_chapter=3)
    save_thesis(doc, mode="123")


def save_thesis(doc, full=True, mode=None):
    if mode is None:
        mode = "full" if full else "345"

    if mode == "full":
        targets = [OUT, OUT_ALT]
        label = "complete thesis"
    elif mode == "123":
        targets = [OUT_CH123, ROOT / "FYP_Chapters_1_2_3.docx"]
        label = "Chapters 1-2-3 only"
    else:
        targets = [OUT_CH345, ROOT / "BAREAI_Chapters_3_4_5.docx"]
        label = "Chapters 3-4-5 only"

    saved = None
    last_err = None
    for path in targets:
        try:
            doc.save(path)
            saved = path
            break
        except PermissionError as e:
            last_err = e
            continue
    if saved is None:
        raise last_err or PermissionError("Could not save document")

    print(f"Saved {label}: {saved}")
    return saved


if __name__ == "__main__":
    import sys

    arg = sys.argv[1] if len(sys.argv) > 1 else ""
    if arg in ("345", "ch345", "--chapters-345"):
        build_chapters_345_only()
    elif arg in ("123", "ch123", "--chapters-123"):
        build_chapters_123_only()
    else:
        build()
