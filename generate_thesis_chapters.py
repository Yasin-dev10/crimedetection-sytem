"""Generate JUST-formatted FYP Chapters 3, 4, and 5 as Word document."""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_margins(section):
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(4.0)
    section.right_margin = Cm(2.0)


def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_paragraph(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=0, space_after=0, line_spacing=2.0, first_line_indent=0):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent:
        pf.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    set_run_font(run, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    if level == 1:
        p = add_paragraph(doc, text, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6)
        for run in p.runs:
            set_run_font(run, size=12, bold=True)
    elif level == 2:
        p = add_paragraph(doc, text, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=10, space_after=4)
        for run in p.runs:
            set_run_font(run, size=12, bold=True)
    else:
        p = add_paragraph(doc, text, bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=4)
        for run in p.runs:
            set_run_font(run, size=12, bold=True, italic=True)
    return p


def add_table(doc, headers, rows, caption=None):
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        cap.paragraph_format.space_after = Pt(6)
        run = cap.add_run(caption)
        set_run_font(run, bold=True, size=12)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            for run in paragraph.runs:
                set_run_font(run, bold=True, size=11)

    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, value in enumerate(row):
            row_cells[c_idx].text = str(value)
            for paragraph in row_cells[c_idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                for run in paragraph.runs:
                    set_run_font(run, size=11)

    doc.add_paragraph()
    return table


def add_figure_placeholder(doc, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("[Insert figure/image here]")
    set_run_font(run, italic=True, size=12)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    cap.paragraph_format.space_after = Pt(12)
    run = cap.add_run(caption)
    set_run_font(run, bold=True, size=12)


def add_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    run = p.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=8)


def build_document():
    doc = Document()
    section = doc.sections[0]
    set_margins(section)
    add_page_number(section)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 2.0
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ==================== CHAPTER III ====================
    add_heading(doc, "CHAPTER III: METHODOLOGY", level=1)

    add_heading(doc, "3.1 Introduction", level=2)
    add_paragraph(doc,
        "This chapter describes the methodology used to develop the Early Warning and Prediction System "
        "for Locust Outbreaks (Locust Prediction System — LPS). It presents the system overview, "
        "architecture, data handling procedures, machine learning approach, requirements, feasibility "
        "analysis, and system design. The methodology follows a structured software engineering and "
        "data science workflow: data acquisition, preprocessing, feature engineering, model development "
        "and evaluation, and integration into a web-based application. The goal is to provide a "
        "reproducible framework for predicting desert locust presence in Somalia using publicly "
        "available environmental data from the Food and Agriculture Organization (FAO)."
    )

    add_heading(doc, "3.2 System Overview", level=2)
    add_paragraph(doc,
        "The Locust Prediction System (LPS) is a web-based early warning platform designed to predict "
        "the likelihood of desert locust presence in a given region based on environmental and temporal "
        "variables. The system accepts inputs such as region, country, year, month, precipitation (PPT), "
        "maximum temperature (TMAX), and soil moisture (SOILMOIS), processes them through a trained "
        "supervised machine learning model, and returns a binary prediction (YES/NO) together with a "
        "confidence score."
    )
    add_paragraph(doc,
        "Unlike traditional locust monitoring methods that rely on manual field surveys and reactive "
        "pesticide spraying, LPS provides proactive, data-driven predictions. Farmers, agricultural "
        "officers, and policymakers can enter current or forecasted environmental conditions and receive "
        "near real-time risk assessments. The system also includes an analytics dashboard, prediction "
        "history, report export (CSV/PDF), and a community blog for sharing outbreak warnings."
    )
    add_paragraph(doc,
        "The system was developed for Somalia and neighbouring regions where locust outbreaks frequently "
        "threaten food security. It prioritizes lightweight algorithms, interpretability, and low "
        "computational cost so that it can operate in resource-constrained environments."
    )

    add_heading(doc, "3.3 System Features", level=2)

    add_heading(doc, "3.3.1 Authentication and Authorization", level=3)
    add_paragraph(doc,
        "The system implements secure user registration and login. Upon successful authentication, users "
        "receive a JSON Web Token (JWT) that restricts access to private data such as personal prediction "
        "history and profile settings."
    )

    add_heading(doc, "3.3.2 Prediction Generation", level=3)
    add_paragraph(doc,
        "Registered users enter environmental parameters through a web form. The backend validates inputs, "
        "applies preprocessing and encoding, invokes the trained model, and returns a locust presence "
        "prediction with probability score."
    )

    add_heading(doc, "3.3.3 Analytics Dashboard", level=3)
    add_paragraph(doc,
        "The dashboard visualizes prediction history using charts and graphs, enabling users to identify "
        "temporal trends, regional risk patterns, and relationships between environmental factors and "
        "locust presence."
    )

    add_heading(doc, "3.3.4 Prediction Feedback", level=3)
    add_paragraph(doc,
        "Users may submit feedback on prediction accuracy in real-world conditions. This supports "
        "continuous model improvement through logged and anonymized data."
    )

    add_heading(doc, "3.3.5 Blog and Content Management", level=3)
    add_paragraph(doc,
        "A built-in blog allows users to create, edit, and delete posts to share locust-related news, "
        "warnings, and community updates."
    )

    add_heading(doc, "3.3.6 Reports and Data Export", level=3)
    add_paragraph(doc,
        "Users can view a detailed history of past predictions and export reports as CSV or PDF, or print "
        "them directly."
    )

    add_heading(doc, "3.3.7 Profile and Account Management", level=3)
    add_paragraph(doc,
        "Users can update profile information, change passwords, and perform account-level actions "
        "including factory reset and account deletion."
    )

    add_heading(doc, "3.4 System Architecture", level=2)
    add_paragraph(doc,
        "The LPS follows a modular three-tier architecture consisting of: (1) Presentation Layer "
        "(Frontend): HTML, CSS, JavaScript, and Bootstrap provide the user interface for data entry, "
        "visualization, and navigation; (2) Application Layer (Backend): Python Flask handles API "
        "requests, authentication, business logic, and model inference; and (3) Data Layer: MySQL "
        "stores user accounts, prediction records, and blog content. The trained machine learning model "
        "is persisted using Joblib."
    )
    add_paragraph(doc,
        "Data flow: User Input → Input Validation → Preprocessing and Encoding → ML Model Prediction "
        "→ Result Display → Optional Logging to Database. The architecture ensures separation of "
        "concerns, scalability, and maintainability. Each component can be updated independently."
    )
    add_figure_placeholder(doc, "Figure 3.1: System Architecture of Locust Prediction System")

    add_heading(doc, "3.5 Data Acquisition & Preprocessing", level=2)

    add_heading(doc, "3.5.1 Dataset Source", level=3)
    add_paragraph(doc,
        "The dataset was obtained from the Food and Agriculture Organization (FAO) and contains "
        "historical locust presence records combined with environmental variables. It comprises 79,529 "
        "records and 8 columns as described in Table 3.1."
    )

    add_table(doc,
        ["Column", "Description"],
        [
            ["REGION", "Sub-national administrative area"],
            ["COUNTRYNAME", "Country name"],
            ["STARTYEAR", "Year of observation (1985–2020)"],
            ["STARTMONTH", "Month of observation (1–12)"],
            ["PPT", "Monthly precipitation (mm)"],
            ["TMAX", "Maximum temperature (°C)"],
            ["SOILMOIS", "Soil moisture index"],
            ["LOCUSTPRESENT", "Target variable (YES/NO)"],
        ],
        caption="Table 3.1: Dataset Column Description"
    )

    add_heading(doc, "3.5.2 Data Preprocessing Steps", level=3)
    add_paragraph(doc,
        "Handling Missing Data: All columns were checked for missing values. No missing values were "
        "found; therefore, no imputation was required."
    )
    add_paragraph(doc,
        "Outlier Detection: Outliers were identified in numerical columns using Interquartile Range (IQR) "
        "analysis. Results: 385 outliers in PPT, 150 in TMAX, and 579 in SOILMOIS. These were handled "
        "to prevent distortion of model training."
    )
    add_paragraph(doc,
        "Encoding Categorical Variables: REGION and COUNTRYNAME were standardized (uppercase, stripped "
        "whitespace) and encoded using target encoding, replacing each category with the mean locust "
        "presence rate for that category."
    )
    add_paragraph(doc,
        "Feature Scaling: Min-max scaling was applied to numerical features (PPT, TMAX, SOILMOIS) to "
        "normalize value ranges."
    )
    add_paragraph(doc,
        "Handling Class Imbalance: The dataset exhibited a class imbalance ratio of approximately 4.7:1 "
        "(absence vs. presence). SMOTE (Synthetic Minority Over-sampling Technique) was applied to "
        "balance the training set and improve recall for the minority class (locust presence)."
    )

    add_heading(doc, "3.6 Feature Engineering / Representation", level=2)
    add_paragraph(doc, "Seven input features were used for model training:")
    add_paragraph(doc,
        "Independent Variables (Features): (1) REGION (target-encoded); (2) COUNTRYNAME "
        "(target-encoded); (3) STARTYEAR; (4) STARTMONTH; (5) PPT (Precipitation); (6) TMAX "
        "(Maximum Temperature); and (7) SOILMOIS (Soil Moisture)."
    )
    add_paragraph(doc,
        "Target Variable: LOCUSTPRESENT — binary classification (1 = YES, 0 = NO)."
    )
    add_paragraph(doc,
        "Data Splitting: The dataset was split into training (80%) and testing (20%) subsets using "
        "train_test_split() with random_state=42 to ensure reproducibility. Domain knowledge guided "
        "feature selection: precipitation and soil moisture support vegetation growth favourable to "
        "locust breeding; temperature affects locust development and activity; temporal features "
        "(year, month) capture seasonal patterns."
    )

    add_heading(doc, "3.7 Model / Algorithm Development", level=2)

    add_heading(doc, "3.7.1 Baseline Model and Model Selection", level=3)
    add_paragraph(doc,
        "Thirteen supervised classification algorithms were evaluated, including Decision Tree, Random "
        "Forest, K-Nearest Neighbors (K-NN), Support Vector Machine (SVM), Logistic Regression, Naïve "
        "Bayes, XGBoost, LightGBM, Gradient Boosting, AdaBoost, CatBoost, Extra Trees Classifier, and "
        "Bagging Classifier. Tree-based ensemble methods were selected as primary candidates due to "
        "their ability to handle non-linear relationships, mixed feature types, and large datasets."
    )

    add_heading(doc, "3.7.2 Model Evaluation", level=3)
    add_paragraph(doc,
        "Models were evaluated using accuracy, precision, recall, F1-score, and AUC-ROC. Hyperparameter "
        "tuning was applied to LightGBM, improving F1-score by 2.32 percentage points while reducing "
        "training time by 65%. The tuned LightGBM classifier was selected for deployment based on test "
        "accuracy of 95.77%, F1-Score of 87.59%, precision of 89.88%, recall of 85.42%, AUC-ROC of "
        "0.982, and inference time of 0.96 seconds."
    )

    add_heading(doc, "3.7.3 Model Deployment Framework", level=3)
    add_paragraph(doc,
        "The trained model was serialized using Joblib (.pkl format) and integrated into a Flask REST "
        "API. The backend exposes endpoints for prediction (/api/predict), saving predictions "
        "(/api/save_prediction), user authentication, and data retrieval. Target encodings for REGION "
        "and COUNTRYNAME are loaded at application startup to ensure consistent preprocessing between "
        "training and inference."
    )

    add_heading(doc, "3.8 Requirements", level=2)

    add_heading(doc, "3.8.1 System Requirements", level=3)
    add_paragraph(doc, "Hardware Requirements", bold=True)

    add_table(doc,
        ["Component", "Specification", "Purpose"],
        [
            ["CPU", "Intel Core i3 (min) / i5+ (recommended)", "Data processing and model inference"],
            ["RAM", "8 GB or higher", "Data preprocessing and model training"],
            ["Storage", "128 GB HDD / 512 GB SSD", "Dataset, model files, and application storage"],
            ["Network", "High-speed broadband", "Online dataset access and web deployment"],
        ],
        caption="Table 3.2: Hardware Requirements"
    )

    add_paragraph(doc, "Software Requirements", bold=True)

    add_table(doc,
        ["Software/Tool", "Version", "Purpose"],
        [
            ["Operating System", "Windows 10+", "Development and deployment platform"],
            ["Python", "3.9+", "Core programming language"],
            ["Flask", "Latest stable", "Web framework and API"],
            ["MySQL", "Latest stable", "Relational database"],
            ["Jupyter Notebook", "Latest", "Data exploration and model training"],
            ["VS Code", "Latest", "Integrated development environment"],
        ],
        caption="Table 3.3: Software Requirements"
    )

    add_table(doc,
        ["Library", "Purpose"],
        [
            ["NumPy", "Numerical computations"],
            ["Pandas", "Data manipulation and analysis"],
            ["Scikit-learn", "Machine learning algorithms"],
            ["Matplotlib / Seaborn", "Data visualization"],
            ["Joblib", "Model persistence"],
            ["imblearn (SMOTE)", "Class imbalance handling"],
        ],
        caption="Table 3.4: Required Libraries"
    )

    add_heading(doc, "3.8.2 User Requirements", level=3)

    add_heading(doc, "3.8.2.1 Functional Requirements", level=3)
    add_table(doc,
        ["ID", "Requirement"],
        [
            ["FR1", "The system shall allow users to register and log in securely."],
            ["FR2", "The system shall accept environmental inputs (region, country, year, month, PPT, TMAX, SOILMOIS)."],
            ["FR3", "The system shall validate all inputs before processing."],
            ["FR4", "The system shall generate locust presence predictions with confidence scores."],
            ["FR5", "The system shall display prediction results clearly on the web interface."],
            ["FR6", "The system shall store prediction history for authenticated users."],
            ["FR7", "The system shall provide an analytics dashboard with charts and graphs."],
            ["FR8", "The system shall allow export of reports in CSV and PDF formats."],
            ["FR9", "The system shall support blog creation, editing, and deletion."],
            ["FR10", "The system shall provide error messages for invalid or incomplete inputs."],
        ],
        caption="Table 3.5: Functional Requirements"
    )

    add_heading(doc, "3.8.2.2 Non-Functional Requirements", level=3)
    add_table(doc,
        ["ID", "Requirement"],
        [
            ["NFR1", "Performance: Predictions shall be generated within 2 seconds of form submission."],
            ["NFR2", "Security: User passwords shall be hashed; API access shall use JWT authentication."],
            ["NFR3", "Usability: The interface shall be intuitive and accessible via modern web browsers."],
            ["NFR4", "Reliability: The system shall remain stable under concurrent user requests."],
            ["NFR5", "Scalability: The architecture shall support additional regions and features."],
            ["NFR6", "Privacy: User data shall not be shared with third parties without consent."],
            ["NFR7", "Maintainability: Code shall follow modular design for easy updates."],
            ["NFR8", "Portability: The system shall run on standard Windows/Linux servers."],
        ],
        caption="Table 3.6: Non-Functional Requirements"
    )

    add_heading(doc, "3.9 Feasibility Study", level=2)

    add_heading(doc, "3.9.1 Technical Feasibility", level=3)
    add_paragraph(doc,
        "The project uses open-source technologies (Python, Flask, MySQL, Scikit-learn) that are well "
        "documented and widely supported. The development team possesses skills in machine learning, "
        "web development, and data preprocessing. The FAO dataset is publicly available and sufficient "
        "for model training."
    )

    add_heading(doc, "3.9.2 Economic Feasibility", level=3)
    add_paragraph(doc,
        "All software tools are free and open-source. Hardware requirements are modest (standard "
        "laptop/desktop). No licensing fees or expensive infrastructure are needed, making the project "
        "affordable for academic and public-sector deployment in low-resource settings."
    )

    add_heading(doc, "3.9.3 Operational Feasibility", level=3)
    add_paragraph(doc,
        "The web interface is designed for non-technical users including farmers and agricultural "
        "officers. Minimal training is required. The system can be accessed from any device with an "
        "internet connection and modern browser."
    )

    add_heading(doc, "3.9.4 Schedule Feasibility", level=3)
    add_paragraph(doc,
        "The project was conducted from January 2025 to August 2025. Major milestones — data collection, "
        "preprocessing, model training, web development, and testing — were completed within this "
        "timeframe, confirming schedule feasibility."
    )

    add_heading(doc, "3.10 System Design", level=2)

    add_heading(doc, "3.10.1 Use Case Diagram", level=3)
    add_paragraph(doc,
        "The system involves two primary actors: Guest User (can view the home page and public blog "
        "content) and Registered User (can register, log in, make predictions, view dashboard, export "
        "reports, manage profile, and create blog posts). Major use cases include Register, Login, Make "
        "Prediction, View Dashboard, Export Report, Manage Blog, and Manage Profile. The Use Case "
        "Diagram is presented in Appendix Figure A.1."
    )

    add_heading(doc, "3.10.2 Database Design", level=3)
    add_paragraph(doc,
        "The MySQL database consists of three main tables: users (id, username, email, password_hash, "
        "security_question, security_answer, created_at), predictions (id, user_id, region, country_name, "
        "start_year, start_month, soil_moisture, tmax, ppt, locust_present, probability, prediction_date), "
        "and blog_posts (id, user_id, title, content, region, country, date, author, tags, image_url). "
        "The database follows Third Normal Form (3NF) to eliminate redundancy. Foreign keys enforce "
        "referential integrity between users and their predictions/blog posts. The Entity-Relationship "
        "Diagram is presented in Appendix Figure A.2."
    )

    doc.add_page_break()

    # ==================== CHAPTER IV ====================
    add_heading(doc, "CHAPTER IV: IMPLEMENTATION AND RESULTS", level=1)

    add_heading(doc, "4.1 Dataset Description", level=2)
    add_paragraph(doc,
        "The dataset used in this study contains 79,529 historical records sourced from the FAO Desert "
        "Locust Information Service. Each record represents a monthly observation for a specific region "
        "and country, with associated environmental conditions and locust presence status."
    )

    add_table(doc,
        ["Attribute", "Type", "Range/Values"],
        [
            ["REGION", "Categorical", "Multiple Somali and regional administrative areas"],
            ["COUNTRYNAME", "Categorical", "Country names"],
            ["STARTYEAR", "Numerical", "1985 – 2020"],
            ["STARTMONTH", "Numerical", "1 – 12"],
            ["PPT", "Numerical", "0 – 500+ mm"],
            ["TMAX", "Numerical", "20 – 45 °C"],
            ["SOILMOIS", "Numerical", "Continuous index"],
            ["LOCUSTPRESENT", "Binary", "YES / NO"],
        ],
        caption="Table 4.1: Dataset Structure and Attributes"
    )

    add_paragraph(doc,
        "Key Statistics (after cleaning): Mean precipitation (PPT) is approximately 50.73 mm; year range "
        "is 1985 to 2020; no missing values were detected; class imbalance ratio is approximately 4.7:1 "
        "(no locust : locust present). Regional distribution shows that regions such as Shebelle, Galbeed, "
        "and Sanaag exhibited higher recorded locust presence compared to Nugaal and Bakool. Notable "
        "spikes in locust activity occurred around 1997, 2007, 2009, 2014, and an unprecedented surge "
        "in 2019–2020."
    )

    add_heading(doc, "4.2 ML Pipeline Implementation", level=2)
    add_paragraph(doc,
        "The machine learning pipeline was implemented in Jupyter Notebook using Python and follows "
        "these stages:"
    )
    add_paragraph(doc,
        "Step 1 — Data Loading and Cleaning: The raw CSV dataset was loaded using Pandas. Missing value "
        "checks confirmed complete data across all columns (Figure 5.1)."
    )
    add_paragraph(doc,
        "Step 2 — Outlier Detection and Handling: IQR-based outlier analysis identified 385 outliers in "
        "PPT, 150 in TMAX, and 579 in SOILMOIS (Figure 5.2). Statistical summaries after cleaning "
        "confirmed data readiness (Figure 5.3)."
    )
    add_paragraph(doc,
        "Step 3 — Exploratory Data Analysis: Visualizations revealed that TMAX in Somalia most "
        "frequently ranges from 32.5°C to 34.0°C (Figure 5.4); locust presence spiked dramatically in "
        "2019–2020 (Figure 5.5); precipitation is generally low with rare high-rainfall events (Figure 5.6); "
        "and regional variation in locust presence exists across Somali regions (Figure 5.7)."
    )
    add_paragraph(doc,
        "Step 4 — Feature Encoding: Target encoding was applied to REGION and COUNTRYNAME columns "
        "(Figure 5.8). LOCUSTPRESENT was binary encoded (YES=1, NO=0)."
    )
    add_paragraph(doc,
        "Step 5 — Data Splitting: 80/20 train-test split with random_state=42 (Figure 5.9)."
    )
    add_paragraph(doc,
        "Step 6 — Model Training and Evaluation: Thirteen classification models were trained and "
        "evaluated (Figures 5.10–5.22)."
    )
    add_paragraph(doc,
        "Step 7 — Model Selection and Serialization: LightGBM (tuned) was selected and saved using "
        "Joblib for Flask deployment."
    )
    add_figure_placeholder(doc, "Figure 4.1: ML Pipeline Implementation")

    add_heading(doc, "4.3 System Design & Deployment", level=2)

    add_heading(doc, "4.3.1 Backend (Flask API)", level=3)
    add_paragraph(doc,
        "The backend was developed using Python Flask and provides RESTful API endpoints for user "
        "registration, authentication (JWT), locust prediction, saving predictions, and data retrieval. "
        "The model and target encodings are loaded at startup. Input data undergoes the same "
        "preprocessing pipeline as training before inference."
    )

    add_table(doc,
        ["Endpoint", "Method", "Function"],
        [
            ["/api/register", "POST", "User registration"],
            ["/api/login", "POST", "User authentication (JWT)"],
            ["/api/predict", "POST", "Locust prediction"],
            ["/api/save_prediction", "POST", "Save prediction to database"],
            ["/api/predictions", "GET/DELETE", "Retrieve/delete predictions"],
        ],
        caption="Table 4.2: Backend API Endpoints"
    )

    add_heading(doc, "4.3.2 Frontend (Web Interface)", level=3)
    add_paragraph(doc,
        "The frontend was built using HTML, CSS, JavaScript, and Bootstrap. Key pages include the Home "
        "Page (Figure 5.23), Login Page (Figure 5.24), Signup Page (Figure 5.25), Dashboard (Figure 5.26), "
        "Prediction Form (Figure 5.27), and Reports Page (Figure 5.28)."
    )

    add_heading(doc, "4.3.3 Integration and Deployment", level=3)
    add_paragraph(doc,
        "The frontend communicates with the Flask backend via AJAX/fetch API calls. Prediction requests "
        "send JSON payloads; responses include prediction label and probability score. The system runs "
        "locally on a development server (Flask) with MySQL as the database backend. It can be deployed "
        "on cloud platforms or institutional servers for production use."
    )

    add_heading(doc, "4.4 Data Visualization (Results Presentation)", level=2)
    add_paragraph(doc,
        "Environmental data visualizations revealed that TMAX in Somalia most frequently ranges from "
        "32.5°C to 34.0°C; precipitation is predominantly low with rare high events; locust presence "
        "showed historical spikes with a dramatic surge in 2019–2020; and regional distribution varies "
        "across Somali regions. Model performance visualizations include confusion matrices, ROC curves, "
        "and feature importance plots highlighting PPT and SOILMOIS as primary drivers. Frontend "
        "screenshots demonstrate the complete user workflow from registration to prediction and report "
        "generation."
    )

    add_heading(doc, "4.5 Model Comparison", level=2)
    add_paragraph(doc,
        "Thirteen models were trained and evaluated on the same test set. Results are summarized in "
        "Table 4.3."
    )

    add_table(doc,
        ["No.", "Model", "Training Accuracy", "Testing Accuracy", "AUC-ROC"],
        [
            ["1", "Decision Tree", "100%", "94.04%", "90.10%"],
            ["2", "Random Forest", "100%", "95.76%", "98.20%"],
            ["3", "K-Nearest Neighbors", "92.37%", "89.53%", "87.17%"],
            ["4", "Support Vector Machine", "92.32%", "92.66%", "98.16%"],
            ["5", "Logistic Regression", "92.32%", "92.66%", "97.32%"],
            ["6", "Naïve Bayes", "92.08%", "92.30%", "96.93%"],
            ["7", "XGBoost", "96.40%", "94.91%", "98.13%"],
            ["8", "LightGBM (Tuned)", "95.02%", "95.77%", "98.20%"],
            ["9", "Gradient Boosting", "93.32%", "93.30%", "97.68%"],
            ["10", "AdaBoost", "92.46%", "92.59%", "91.50%"],
            ["11", "CatBoost", "96.73%", "94.46%", "95.48%"],
            ["12", "Extra Trees", "100%", "95.16%", "94.80%"],
            ["13", "Bagging Classifier", "99.74%", "95.17%", "95.46%"],
        ],
        caption="Table 4.3: Model Performance Comparison"
    )

    add_table(doc,
        ["Metric", "Value"],
        [
            ["Test Accuracy", "95.77%"],
            ["Precision", "89.88%"],
            ["Recall", "85.42%"],
            ["F1-Score", "87.59%"],
            ["AUC-ROC", "0.982"],
            ["Inference Time", "0.96 seconds"],
        ],
        caption="Table 4.4: Detailed Metrics for Selected Model (LightGBM Tuned)"
    )

    add_paragraph(doc,
        "Analysis: Tree-based ensembles (Random Forest, LightGBM, XGBoost) consistently outperformed "
        "simpler models. LightGBM achieved the best balance of accuracy, F1-score, and inference speed. "
        "Random Forest matched LightGBM in AUC-ROC (0.982) but had slightly lower recall. SVM achieved "
        "comparable accuracy (95.78%) but required significantly longer training time (~1481 seconds vs. "
        "~5.88 seconds for LightGBM). K-NN showed the lowest performance (89.53% test accuracy). "
        "Hyperparameter tuning improved LightGBM F1-score by 2.32 percentage points."
    )

    doc.add_page_break()

    # ==================== CHAPTER V ====================
    add_heading(doc, "CHAPTER V: DISCUSSION AND CONCLUSION", level=1)

    add_heading(doc, "5.1 Introduction", level=2)
    add_paragraph(doc,
        "This chapter interprets the results presented in Chapter IV, discusses the contributions of this "
        "research, identifies challenges and limitations, and outlines recommendations for future work. "
        "The discussion connects findings back to the research objectives and questions stated in Chapter I."
    )

    add_heading(doc, "5.2 Interpretation of Results", level=2)
    add_paragraph(doc,
        "The results demonstrate that machine learning can effectively predict desert locust presence "
        "using a minimal set of environmental variables. The tuned LightGBM model achieved 95.77% test "
        "accuracy with an F1-score of 87.59% and AUC-ROC of 0.982, directly addressing Research "
        "Question 2 regarding model evaluation."
    )
    add_paragraph(doc,
        "Research Objective 1 — implementing an ML model for locust prediction — was achieved through "
        "training and comparison of thirteen algorithms on FAO data. LightGBM was selected for its "
        "superior balance of predictive performance and computational efficiency."
    )
    add_paragraph(doc,
        "Research Objective 2 — evaluating model accuracy and robustness — was fulfilled through "
        "comprehensive metrics (accuracy, precision, recall, F1, AUC-ROC) and an 80/20 train-test "
        "split with SMOTE handling of class imbalance."
    )
    add_paragraph(doc,
        "Research Objective 3 — developing a user-friendly web application — was achieved through "
        "LocustHub, which provides prediction forms, dashboards, reports, and a community blog."
    )
    add_paragraph(doc,
        "Feature importance analysis confirmed that precipitation (PPT) and soil moisture (SOILMOIS) are "
        "the primary drivers of locust presence predictions, aligning with established entomological "
        "knowledge. The 2019–2020 locust surge visible in the data corresponds to the well-documented "
        "East Africa desert locust crisis (FAO, 2020)."
    )

    add_heading(doc, "5.3 Contribution to Knowledge", level=2)
    add_paragraph(doc,
        "This research makes the following contributions: (1) Operational Accessibility — unlike complex "
        "satellite-based or drone-based systems, LPS uses readily available environmental data; "
        "(2) Comprehensive Model Comparison — this study trained and compared thirteen classification "
        "algorithms, providing a thorough benchmark for locust prediction; (3) User-Centric Design — "
        "LPS includes a complete web application bridging the gap between ML research and practical "
        "deployment; (4) Class Imbalance Handling — successful application of SMOTE achieving 85.42% "
        "recall for the critical minority class; and (5) Regional Focus — the study specifically targets "
        "Somalia, a region severely affected by locust outbreaks but underrepresented in existing "
        "ML-based locust prediction literature."
    )

    add_heading(doc, "5.4 Challenges and Limitations", level=2)
    add_paragraph(doc,
        "Several challenges were encountered during this study. Data quality and availability: the dataset "
        "relies on historical FAO records which may contain reporting inconsistencies; regionally "
        "disaggregated real-time data for Somalia is limited; additional variables such as humidity, wind "
        "speed, or NDVI are not included. Class imbalance: the 4.7:1 ratio required SMOTE, which "
        "introduces synthetic samples that may not fully represent real-world outbreak patterns."
    )
    add_paragraph(doc,
        "Model generalization: the model was trained primarily on Somalia and neighbouring region data; "
        "generalization to other continents requires further validation; some models showed signs of "
        "overfitting with 100% training accuracy. Computational constraints: SVM required approximately "
        "1481 seconds for training. Deployment limitations: the system runs on a local development server; "
        "no mobile application exists; real-time data integration is not yet implemented."
    )

    add_heading(doc, "5.5 Future Work", level=2)
    add_paragraph(doc,
        "Based on the findings and limitations identified, the following future directions are recommended: "
        "(1) Geographical Expansion — extend the system to cover additional East and West African "
        "countries; (2) Satellite Data Integration — incorporate NDVI and land surface temperature; "
        "(3) Real-Time Prediction — integrate live weather API feeds; (4) Mobile Application — develop "
        "Android/iOS applications for field officers and farmers; (5) Multilingual Support — add Somali, "
        "Arabic, and other widely spoken languages; (6) Advanced Models — explore deep learning "
        "architectures (LSTM, CNN) for temporal and spatial pattern recognition; (7) Government "
        "Collaboration — partner with agricultural ministries and meteorological departments; and "
        "(8) Feedback Loop — implement automated model retraining using user-submitted prediction feedback."
    )

    add_heading(doc, "5.6 Final Conclusion", level=2)
    add_paragraph(doc,
        "This study successfully developed an Early Warning and Prediction System for Locust Outbreaks "
        "using machine learning techniques. The system demonstrates that accurate locust presence "
        "prediction (95.77% accuracy) is achievable using publicly available environmental data and "
        "lightweight algorithms, without requiring expensive satellite infrastructure or drone fleets."
    )
    add_paragraph(doc,
        "The LocustHub web application transforms sophisticated machine learning capabilities into an "
        "accessible tool for farmers, agricultural officers, and policymakers in Somalia. By providing "
        "early warnings based on precipitation, temperature, and soil moisture, the system enables "
        "proactive interventions that can reduce crop losses and strengthen food security."
    )
    add_paragraph(doc,
        "The research confirms that data-driven approaches offer a cost-effective, scalable alternative "
        "to traditional reactive locust management methods. With further enhancements — including "
        "satellite data integration, mobile access, and real-time prediction — systems like LPS can "
        "play a significant role in mitigating the devastating impact of locust outbreaks on agriculture "
        "and livelihoods in vulnerable regions."
    )

    output_path = r"C:\Users\pc\Downloads\BAREAI-main-main (4)\BAREAI-main-main\FYP_Chapters_3_4_5.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    build_document()
