from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "3 chapters" / "chapter-one-two-three_REVISED.docx"
OUTPUT = ROOT / "3 chapters" / "chapter-one-two-three_LINKS_FIXED.docx"


def set_run_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")


def append_field(paragraph, instruction, visible_text):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instruction)
    fld.set(qn("w:dirty"), "true")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_pr.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "24")
    r_pr.append(size)
    run.append(r_pr)
    text = OxmlElement("w:t")
    text.text = visible_text
    run.append(text)
    fld.append(run)
    paragraph._p.append(fld)
    return fld


def insert_paragraph_before_element(element, parent, style=None):
    new_p = OxmlElement("w:p")
    element.addprevious(new_p)
    p = Paragraph(new_p, parent)
    if style:
        p.style = style
    return p


def create_caption_before_table(table, label, chapter, title, reset=False):
    p = insert_paragraph_before_element(table._tbl, table._parent, "Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    prefix = p.add_run(f"{label} {chapter}.")
    set_run_font(prefix, 11, True)
    instruction = f"SEQ {label} \\* ARABIC"
    if reset:
        instruction += " \\r 1"
    append_field(p, instruction, "1")
    suffix = p.add_run(f" - {title}")
    set_run_font(suffix, 11, False)
    return p


def replace_figure_caption(paragraph, number, title, reset=False):
    paragraph.clear()
    paragraph.style = "Caption"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.keep_with_next = False
    prefix = paragraph.add_run("Figure 3.")
    set_run_font(prefix, 11, True)
    instruction = "SEQ Figure \\* ARABIC"
    if reset:
        instruction += " \\r 1"
    append_field(paragraph, instruction, str(number))
    suffix = paragraph.add_run(f" - {title}")
    set_run_font(suffix, 11, False)


def make_front_heading(text):
    p = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    p_pr.append(jc)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "240")
    p_pr.append(spacing)
    p.append(p_pr)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    bold = OxmlElement("w:b")
    r_pr.append(bold)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_pr.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "28")
    r_pr.append(size)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    p.append(run)
    return p


def make_page_break():
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    return p


def make_list_field(label):
    p = OxmlElement("w:p")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f'TOC \\h \\z \\c "{label}"')
    fld.set(qn("w:dirty"), "true")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = f"Right-click and choose Update Field to populate the List of {label}s."
    r.append(t)
    fld.append(r)
    p.append(fld)
    return p


doc = Document(SOURCE)

# Remove stray cover-page placeholders and the obsolete empty List-of-Figures field.
for p in list(doc.paragraphs):
    instructions = [n.text or "" for n in p._p.xpath(".//w:instrText")]
    instructions += [
        n.get(qn("w:instr")) or "" for n in p._p.xpath(".//w:fldSimple")
    ]
    joined_instructions = " ".join(instructions)
    if "TOC" in joined_instructions and '\\c "Figure"' in joined_instructions:
        parent = p._p.getparent()
        if parent is not None:
            parent.remove(p._p)
        continue
    if p.text.strip().lower() == "list of tables":
        p.clear()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

# Correct caption misuse and normalize visible caption encoding.
for p in doc.paragraphs:
    text = " ".join(p.text.split())
    if text == "3.4 System Features":
        p.style = "Heading 1"
    elif text.startswith("Figure 3.1"):
        replace_figure_caption(p, 1, "System Architecture of BAREAI", reset=True)
    elif text.startswith("Figure 3.2"):
        replace_figure_caption(p, 2, "ML Training and Inference Pipeline")

# Create proper Caption paragraphs for all seven tables and clean embedded titles.
table_specs = [
    ("Table", "2", "Summary of Related Studies", True, "Author & Year"),
    ("Table", "3", "Dataset Schema", True, "Attribute"),
    ("Table", "3", "Model Comparison Criteria", False, "Criterion"),
    ("Table", "3", "Development Tools", False, "Tool"),
    ("Table", "3", "Hardware Requirements", False, "Component"),
    ("Table", "3", "Software Requirements", False, "Software"),
    ("Table", "3", "Required Libraries and Frameworks", False, "Library / Framework"),
]
for table, spec in zip(doc.tables, table_specs):
    label, chapter, title, reset, first_header = spec
    create_caption_before_table(table, label, chapter, title, reset)
    first_cell = table.rows[0].cells[0]
    first_p = first_cell.paragraphs[0]
    first_p.clear()
    run = first_p.add_run(first_header)
    set_run_font(run, 10, True)

# Update narrative table references to the corrected chapter-based sequence.
for p in doc.paragraphs:
    if p.text.strip().startswith("Table 3.1 specifies minimum"):
        for run in p.runs:
            run.text = run.text.replace("Table 3.1", "Table 3.4")
    if p.text.strip().startswith("Table 3.3 lists the primary"):
        for run in p.runs:
            run.text = run.text.replace("Table 3.3", "Table 3.6")

# Add university-style front-matter lists immediately before Chapter I.
chapter_one = next(
    p for p in doc.paragraphs if p.text.strip().upper() == "CHAPTER I: INTRODUCTION"
)
anchor = chapter_one._p
for element in [
    make_page_break(),
    make_front_heading("LIST OF FIGURES"),
    make_list_field("Figure"),
    make_page_break(),
    make_front_heading("LIST OF TABLES"),
    make_list_field("Table"),
]:
    anchor.addprevious(element)

# Ensure Word refreshes every generated field when the document opens.
settings = doc.settings._element
update_fields = settings.find(qn("w:updateFields"))
if update_fields is None:
    update_fields = OxmlElement("w:updateFields")
    settings.append(update_fields)
update_fields.set(qn("w:val"), "true")

doc.core_properties.comments = (
    "TOC, List of Figures, and List of Tables repaired with clickable Word fields."
)
doc.save(OUTPUT)
print(OUTPUT)
