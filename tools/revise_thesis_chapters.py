from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "3 chapters" / "chapter-one , twoo & three_backup.docx"
OUTPUT = ROOT / "3 chapters" / "chapter-one-two-three_REVISED.docx"
NUM_ID = None


def find_paragraph(doc, prefix):
    for p in doc.paragraphs:
        if " ".join(p.text.split()).lower().startswith(prefix.lower()):
            return p
    raise ValueError(f"Paragraph not found: {prefix}")


def find_chapter_heading(doc, roman):
    target = f"CHAPTER {roman}"
    for p in doc.paragraphs:
        text = " ".join(p.text.split()).upper()
        if p.style.name.startswith("Heading") and text.startswith(target):
            return p
    raise ValueError(f"Chapter heading not found: {roman}")


def set_text(p, text, style=None):
    p.clear()
    p.add_run(text)
    if style == "List Number":
        style = "Normal"
        p_pr = p._p.get_or_add_pPr()
        num_pr = p_pr.get_or_add_numPr()
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), str(NUM_ID))
        num_pr.append(ilvl)
        num_pr.append(num_id)
    if style:
        p.style = style


def insert_before(anchor, text, style="Normal"):
    p = anchor.insert_paragraph_before()
    set_text(p, text, style)
    return p


def insert_after(anchor, text, style="Normal"):
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    p = Paragraph(new_p, anchor._parent)
    set_text(p, text, style)
    return p


def clear_between(doc, start_prefix, end_prefix):
    start = find_paragraph(doc, start_prefix)
    end = find_paragraph(doc, end_prefix)
    active = False
    for p in doc.paragraphs:
        if p._p is start._p:
            active = True
            continue
        if p._p is end._p:
            break
        if active:
            set_text(p, "")
    return start, end


def clear_between_paragraphs(doc, start, end):
    active = False
    for p in doc.paragraphs:
        if p._p is start._p:
            active = True
            continue
        if p._p is end._p:
            break
        if active:
            set_text(p, "")


def remove_safe_empty_paragraphs_between(doc, start, end):
    active = False
    for p in list(doc.paragraphs):
        if p._p is start._p:
            active = True
            continue
        if p._p is end._p:
            break
        if not active or p.text.strip():
            continue
        if p._p.xpath(".//w:drawing | .//w:pict | .//w:sectPr | .//w:bookmarkStart"):
            continue
        parent = p._p.getparent()
        if parent is not None:
            parent.remove(p._p)


def add_block_before(anchor, blocks):
    for style, text in blocks:
        insert_before(anchor, text, style)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def create_numbering_definition(document):
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(tabs)
    p_pr.append(ind)
    lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def normalize_mojibake(text):
    replacements = {
        "â€™": "'",
        "â€œ": '"',
        "â€": '"',
        "â€“": "-",
        "â€”": "-",
        "â€‘": "-",
        "â€\x90": "-",
        "â†’": "to",
        "crimerelated": "crime-related",
        "artificialintelligence": "artificial intelligence",
        "lowresource": "low-resource",
        "MachineLearning": "Machine Learning",
        "CybercrimeDetection": "Cybercrime Detection",
        "EGovernment": "E-Government",
    }
    for bad, good in replacements.items():
        text = text.replace(bad, good)
    text = re.sub(r"\s+([,.;:?!])", r"\1", text)
    text = re.sub(r"\.(?=\()", ". ", text)
    text = re.sub(r"\)(?=[A-Z])", ") ", text)
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()


doc = Document(SOURCE)
NUM_ID = create_numbering_definition(doc)

# Chapter I is rewritten for logical alignment among problem, objectives, questions, and scope.
first_ch2 = find_chapter_heading(doc, "II")
ch1_start = find_paragraph(doc, "1.1 Background")
for p in list(doc.paragraphs):
    if p._p is ch1_start._p:
        break
chapter_one_title = insert_before(ch1_start, "CHAPTER I: INTRODUCTION", "Heading 1")
chapter_one_title.paragraph_format.page_break_before = True

ch2_title = first_ch2
ch2_title.paragraph_format.page_break_before = True

clear_between_paragraphs(doc, ch1_start, first_ch2)
ch2_anchor = first_ch2
set_text(ch1_start, "1.1 Background of the Study", "Heading 1")
chapter1_blocks = [
    ("Normal", "Crime-related information is increasingly communicated through online news, social media, email, and web-based reporting channels. Although these sources can support timely situational awareness, the reports are commonly written as unstructured text and must be interpreted before they can be routed, prioritised, or investigated. Manual review becomes slow and inconsistent when the volume of reports grows."),
    ("Normal", "Natural Language Processing (NLP) provides computational methods for converting unstructured language into features that machine-learning algorithms can analyse. In crime-report processing, NLP can support automatic classification, information extraction, and decision support. Previous studies show that supervised learning and transformer-based approaches can classify crime narratives and extract investigative information, but performance depends on the language, label quality, domain, and availability of representative training data (Carnaz et al., 2021; Hariguna & Ruangkanjanases, 2023; Park et al., 2024)."),
    ("Normal", "The challenge is more pronounced for Somali-language content. Somali is comparatively under-resourced in mainstream NLP research, and online reports often contain spelling variation, borrowed English terms, informal phrasing, and code-mixed text. Models developed primarily for high-resource languages may therefore transfer poorly without local preprocessing, representative data, and domain-specific evaluation."),
    ("Normal", "This study develops BAREAI, an AI-assisted platform for automatically classifying Somali and mixed-language text as crime-related or not crime-related. The platform combines a supervised text classifier with Somali-language preprocessing and operational features such as location and keyword extraction, blacklist monitoring, notifications, and investigation case management. The intention is to support analysts' decisions; the system does not determine guilt or replace professional investigation."),
    ("Heading 1", "1.2 Problem Statement"),
    ("Normal", "Security organisations need to review large volumes of textual information quickly so that potentially urgent reports can be identified and routed for further assessment. In the current manual workflow, an officer or analyst reads each report, decides whether it is crime-related, and records or forwards it. This process is time-consuming, difficult to scale, and vulnerable to inconsistent judgement, fatigue, and delayed response."),
    ("Normal", "Existing automated text-classification solutions are commonly developed for English or other comparatively well-resourced languages and may not address Somali linguistic variation or the operational workflow required by local investigators. A further limitation is that classification alone does not connect the prediction to monitoring, alerts, and case-management activities. Consequently, there is a need for an integrated and empirically evaluated system that can process Somali-oriented text, distinguish crime-related from non-crime-related content, present transparent confidence information, and support human follow-up."),
    ("Heading 1", "1.3 Research Objectives"),
    ("Heading 2", "1.3.1 General Objective"),
    ("Normal", "To design, implement, and evaluate an NLP-based system that automatically classifies Somali-oriented text reports as crime-related or not crime-related and supports subsequent monitoring and investigation workflows."),
    ("Heading 2", "1.3.2 Specific Objectives"),
    ("Normal", "SO1: To collect, clean, label, and prepare a representative dataset of crime-related and non-crime-related text."),
    ("Normal", "SO2: To transform the prepared text into machine-readable features using an NLP preprocessing and TF-IDF representation pipeline."),
    ("Normal", "SO3: To train and compare supervised machine-learning classifiers and select a production model using accuracy, precision, recall, F1-score, and crime-class recall."),
    ("Normal", "SO4: To integrate the selected classifier into a web platform that provides text analysis, monitoring, alerts, and investigation case management."),
    ("Normal", "SO5: To evaluate the technical performance and practical limitations of the resulting system."),
    ("Heading 1", "1.4 Research Questions"),
    ("Normal", "RQ1: How can Somali-oriented crime text be collected, labelled, and preprocessed for reliable binary classification?"),
    ("Normal", "RQ2: Which TF-IDF configuration and supervised machine-learning classifier provide the best balance of overall F1-score and crime-class recall?"),
    ("Normal", "RQ3: How can the selected classifier be integrated into a usable web-based monitoring and investigation workflow?"),
    ("Normal", "RQ4: What technical limitations and error patterns remain after implementation and evaluation?"),
    ("Heading 1", "1.5 Motivation of the Study"),
    ("Normal", "The study is motivated by the growing amount of Somali-language information published through digital channels and the limited availability of locally adapted tools for analysing it. Faster initial screening can help analysts focus attention on relevant reports while preserving human review for sensitive decisions. The project also provides a practical opportunity to develop and evaluate NLP methods for a low-resource language in a socially important domain."),
    ("Heading 1", "1.6 Significance of the Study"),
    ("Normal", "For security organisations, the system offers a consistent first-stage screening mechanism that can reduce repetitive manual sorting and make potentially relevant content easier to locate. For investigators, integrated alerts, history, blacklist matching, and case-management functions connect model output to a documented follow-up workflow. For researchers and developers, the study contributes an applied Somali-oriented NLP pipeline, an evaluation of conventional supervised models, and implementation lessons for deploying language technology in a low-resource setting."),
    ("Heading 1", "1.7 Scope of the Study"),
    ("Heading 2", "1.7.1 Content and Functional Scope"),
    ("Normal", "The study is limited to binary text classification: crime-related and not crime-related. It covers text preprocessing, TF-IDF feature extraction, supervised model comparison, a Flask inference service, an Express/MongoDB backend, and a React interface. Operational functions include manual text analysis, monitored-source ingestion, location and keyword extraction, blacklist matching, notifications, reporting, and investigation case management."),
    ("Heading 2", "1.7.2 Language and Data Scope"),
    ("Normal", "The dataset contains Somali-oriented text and may include English terms or code-mixed content. The model is evaluated only against the labelled dataset used in this project. Its performance should not be assumed to generalise to every dialect, platform, crime category, jurisdiction, or future language pattern without additional validation and retraining."),
    ("Heading 2", "1.7.3 Study Limitations"),
    ("Normal", "The platform identifies linguistic patterns associated with crime-related reporting; it does not verify whether an allegation is true, predict that a person will commit a crime, or replace legal and investigative judgement. Predictions may contain false positives and false negatives. Access to external platforms, including Facebook, also depends on valid APIs, permissions, and platform policies."),
    ("Heading 1", "1.8 Organization of the Study"),
    ("Normal", "Chapter I introduces the study, defines the problem, presents the objectives and research questions, and explains the motivation, significance, and scope. Chapter II critically reviews literature on crime-related text, social-media monitoring, conventional and machine-learning classification, low-resource language challenges, and related systems before identifying the research gap. Chapter III describes the system architecture, dataset preparation, feature representation, model development and evaluation, implementation approach, development environment, and requirements. Chapter IV presents system analysis, design, implementation, and results. Chapter V summarises the findings, concludes the study, and provides recommendations for future work."),
]
add_block_before(ch2_anchor, chapter1_blocks)

# Strengthen the synthesis at the start and end of Chapter II while retaining the cited thematic review.
ch2_intro, ch22 = clear_between(doc, "2.1 Introduction", "2.2 Impact")
set_text(ch2_intro, "2.1 Introduction", "Heading 1")
add_block_before(ch22, [
    ("Normal", "This chapter reviews scholarship relevant to automatic classification of crime-related text. The review is organised around five themes: the effects and reporting of crime on social media, conventional classification practices, cybercrime detection, technical challenges in crime-text classification, and closely related machine-learning studies. The purpose is not merely to list previous work, but to compare their data, methods, evaluation choices, and limitations in relation to Somali-language deployment."),
    ("Normal", "Existing studies demonstrate that NLP and machine learning can support crime-report triage, crime-news classification, event extraction, and cybercrime complaint processing (Carnaz et al., 2021; Hariguna & Ruangkanjanases, 2023; Rani et al., 2024). However, the literature also identifies recurring constraints: class imbalance, limited labelled data, domain shift, multilingual and code-mixed language, weak interpretability, and the operational consequences of false negatives."),
    ("Normal", "The review therefore distinguishes between prediction performance in an experimental dataset and usefulness in an operational environment. A model with high overall accuracy may still be unsuitable if it misses a substantial proportion of crime-related cases, leaks information from the test set, or is not integrated with a human review process. These considerations guide the methodology presented in Chapter III."),
])

related, proposed = clear_between(doc, "2.6 Related Work", "2.7 Proposed System")
set_text(related, "2.6 Related Work", "Heading 1")
add_block_before(proposed, [
    ("Normal", "Hariguna and Ruangkanjanases (2023) proposed an adaptive decision-support model for automated analysis and classification of crime reports. Their work demonstrates the value of structuring incoming narratives for e-government decision making, but it does not focus on Somali-language data or an end-to-end investigation workflow."),
    ("Normal", "Carnaz et al. (2021) developed an annotated Portuguese corpus for crime-related NLP and machine-learning tasks. The study is particularly relevant because it shows that language-specific labelled resources are essential outside English; however, Portuguese linguistic resources and institutional contexts cannot be transferred directly to Somali text."),
    ("Normal", "Ali et al. (2023) and Rani et al. (2024) investigated transformer-based approaches for crime-news and code-mixed cybercrime classification. Their findings support contextual language modelling for complex narratives, but transformer deployment may require greater computational resources and larger, carefully labelled datasets than conventional TF-IDF models."),
    ("Normal", "Mussiraliyeva and Baispay (2024) compared machine-learning approaches for crime-related textual data using multiple evaluation measures. Their comparative design informs this study's model-selection procedure. Park et al. (2024), meanwhile, showed that transformer models can extract key investigative information from legal documents, although information extraction differs from the binary classification objective of BAREAI."),
    ("Normal", "Across these studies, three patterns are clear. First, language and domain strongly influence performance. Second, transparent train/validation/test separation and class-sensitive metrics are necessary for credible evaluation. Third, most published classifiers emphasise model output and provide limited integration with monitoring, alerts, and case management. These observations define the gap addressed by the present study."),
])
set_text(proposed, "2.7 Proposed System", "Heading 1")
gaps = find_paragraph(doc, "2.8 Gaps")
clear_between_paragraphs(doc, proposed, gaps)
insert_after(proposed, "BAREAI is a web-based crime-intelligence support platform for Somali-oriented text. It accepts user-entered or monitored text, applies the same preprocessing and TF-IDF representation used during training, and returns a binary crime-related decision with confidence information. The selected production classifier is exposed through a Flask service and integrated with an Express/MongoDB backend and React frontend. The wider platform provides source monitoring, Somali keyword and location enrichment, blacklist matching, notifications, report generation, and investigation case management. All model outputs remain subject to human review.")
set_text(gaps, "2.8 Research Gap and Study Contribution", "Heading 1")
gap_body = next(p for p in doc.paragraphs if p._p.getprevious() is gaps._p)
set_text(gap_body, "The reviewed literature provides strong evidence for automated crime-text analysis, but it offers limited support for Somali-language and code-mixed content and rarely connects classification to an operational investigation workflow. This study addresses that gap by: (1) preparing and evaluating a Somali-oriented labelled dataset; (2) preventing data leakage through a train-only TF-IDF fit; (3) comparing multiple supervised classifiers with emphasis on F1-score and crime-class recall; (4) deploying the selected model through a reproducible service architecture; and (5) integrating classification with monitoring, alerts, blacklist checks, reporting, and case management. The contribution is an applied decision-support system, not an autonomous policing or guilt-determination tool.")
ch3 = find_chapter_heading(doc, "III")
insert_before(ch3, "2.9 Chapter Summary", "Heading 1")
insert_before(ch3, "This chapter reviewed the social and technical context of crime-related text classification, compared relevant NLP and machine-learning studies, and identified the need for a locally adapted, carefully evaluated, and operationally integrated Somali-language solution. The next chapter explains the methodology used to prepare the data, develop and select the model, and implement the BAREAI platform.", "Normal")
ch3.paragraph_format.page_break_before = True
remove_safe_empty_paragraphs_between(doc, chapter_one_title, first_ch2)
remove_safe_empty_paragraphs_between(doc, first_ch2, ch3)

# Correct methodology statements so they match the actual repository pipeline.
replacements = {
    "3.5.2.6 Train-Test Split": "3.5.2.6 Train-Validation-Test Split",
    "The prepared dataset is divided into training": "The prepared dataset is divided using stratified sampling into 70% training, 15% validation, and 15% held-out test partitions. TF-IDF is fitted only on the training partition and then used to transform validation and test data. The validation set supports feature and model selection, while the untouched test set provides the final unbiased evaluation.",
    "If class imbalance is detected": "Class distribution is preserved through stratified splitting. Models that support class weighting, including Logistic Regression, Linear SVM, Decision Tree, and Random Forest, are configured with balanced class weights so that minority-class errors receive appropriate attention. The held-out test distribution is not synthetically altered.",
    "Class distribution was analyzed": "Class distribution was examined during exploratory data analysis. Stratified splitting preserves the observed label proportions across training, validation, and test partitions, while supported classifiers use balanced class weights. No synthetic samples are introduced into the held-out test data.",
    "All models were trained on the 80% training partition": "The dataset was divided using stratified sampling into 70% training, 15% validation, and 15% held-out test partitions. TF-IDF was fitted only on the training text to prevent data leakage. Baseline models and feature configurations were compared on validation data; selected models were then tuned through three-fold stratified cross-validation on the combined training and validation partitions. Final performance was measured once on the untouched test set. Evaluation included accuracy, weighted precision, weighted recall, weighted F1-score, crime-class recall, confusion matrix, and ROC-AUC.",
    "The model achieving the highest F1-score": "The tuned Random Forest achieved the strongest held-out performance in the final experiment: 92.3% accuracy, 92.3% weighted F1-score, and 95.9% recall for the crime-related class. Logistic Regression and Linear SVM achieved weighted F1-scores of 91.5% and 91.3%, respectively. Random Forest was therefore selected for deployment, and the classifier, fitted TF-IDF vectorizer, and model metadata were serialised as crime_model.pkl, vectorizer.pkl, and model_meta.pkl.",
    "The BAREAI implementation fine-tunes": "BERT was reviewed as a possible contextual benchmark, but it was not the deployed model in the completed implementation. The production pipeline uses TF-IDF and conventional supervised classifiers because they provided strong measured performance with substantially lower training and inference requirements. Future work may evaluate multilingual or African-language transformer models under the same leakage-free split and held-out test protocol.",
}
for prefix, new_text in replacements.items():
    try:
        p = find_paragraph(doc, prefix)
        if prefix.startswith("3.5.2.6"):
            set_text(p, new_text, "Heading 3")
        else:
            set_text(p, new_text)
    except ValueError:
        pass

bert_h = find_paragraph(doc, "3.5.3.3 BERT")
set_text(bert_h, "3.5.3.3 Transformer Model Consideration (BERT)", "Heading 3")
feature_h = find_paragraph(doc, "3.4.1 Authentication")
insert_before(feature_h, "3.4 System Features", "Heading 1")
training_h = find_paragraph(doc, "3.5.4 Model Training")
training_h.style = "Heading 2"
hybrid_h = find_paragraph(doc, "3.5.5 Hybrid")
set_text(hybrid_h, "3.5.5 Hybrid Enrichment and Inference", "Heading 2")
implementation_h = find_paragraph(doc, "3.5.6 System Implementation")
implementation_h.style = "Heading 2"
try:
    requirement_p = find_paragraph(doc, "Table 3.1 specifies")
    set_text(requirement_p, "Table 3.1 specifies minimum and recommended hardware configurations. The deployed TF-IDF and Random Forest pipeline can operate on a multi-core CPU; additional memory improves concurrent analysis and monitoring. GPU acceleration is not required for the production model, although it would be useful for future transformer experiments.")
except ValueError:
    pass

# Standardise headings and repair visible encoding/spacing problems.
heading_fixes = {
    "CHAPTER II:LITERATURE REVIEW": "CHAPTER II: LITERATURE REVIEW",
    "2.2 Impact of crimes on social media": "2.2 Crime and Social Media",
    "2.3 Traditional Classification techniques in crimes": "2.3 Traditional Crime-Text Classification Approaches",
    "2.4 Cybercrime detection in social media": "2.4 Cybercrime Detection on Social Media",
}
for p in doc.paragraphs:
    clean = normalize_mojibake(p.text)
    if clean in heading_fixes:
        clean = heading_fixes[clean]
    if clean != p.text.strip():
        set_text(p, clean, p.style)

# Academic page setup and style system.
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.first_line_indent = Cm(1.27)

for name, size, before, after in [
    ("Heading 1", 14, 12, 6),
    ("Heading 2", 12, 10, 4),
    ("Heading 3", 12, 8, 3),
]:
    st = styles[name]
    st.font.name = "Times New Roman"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor(0, 0, 0)
    st._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True
    st.paragraph_format.first_line_indent = Cm(0)

for p in doc.paragraphs:
    text = p.text.strip()
    if not text:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        continue
    for run in p.runs:
        run.font.name = "Times New Roman"
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Times New Roman")
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Times New Roman")
    if p.style.name == "Normal" and not re.match(r"^(AUTOMATIC|YAASIIN|NAIMA|NASTEHA|NAJMA|SUBMISSION|PARTIAL|DEGREE|COMPUTER|JAHMURIYA|FACULTY|AUGUST)", text):
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.widow_control = True
    if text.upper().startswith("CHAPTER "):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.page_break_before = True
        p.paragraph_format.keep_with_next = True

for table in doc.tables:
    table.autofit = True
    if table.rows:
        set_repeat_table_header(table.rows[0])
        for cell in table.rows[0].cells:
            for run in cell.paragraphs[0].runs:
                run.bold = True
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
            if "transformers" in cell.text.lower():
                for p in cell.paragraphs:
                    if "transformers" in p.text.lower():
                        set_text(p, "transformers (optional future evaluation)")
                        for run in p.runs:
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(10)

references_heading = find_paragraph(doc, "REFERENCES")
references_heading.paragraph_format.page_break_before = True

# Consolidate line-broken references into readable APA-style hanging paragraphs.
all_paragraphs = doc.paragraphs
ref_index = next(i for i, p in enumerate(all_paragraphs) if p._p is references_heading._p)
source_refs = [p for p in all_paragraphs[ref_index + 1 :] if p.text.strip()]
groups = []
for p in source_refs:
    text = normalize_mojibake(p.text)
    is_author_start = bool(
        re.match(r"^[A-ZÀ-ÖØ-Ý][A-Za-zÀ-ÖØ-öø-ÿ'’\-]+,\s+[A-Z]", text)
    )
    if is_author_start or not groups:
        groups.append([text])
    else:
        groups[-1].append(text)
for p in source_refs:
    parent = p._p.getparent()
    if parent is not None:
        parent.remove(p._p)
anchor = references_heading
for fragments in groups:
    p = insert_after(anchor, " ".join(fragments), "Normal")
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
    anchor = p

doc.core_properties.title = "Automatic Classification of Crime-Related Text Reports Using NLP"
doc.core_properties.subject = "Revised Chapters I-III"
doc.core_properties.comments = "Revised in accordance with the AI-based FYP guideline and thesis benchmark."
doc.save(OUTPUT)
print(OUTPUT)
