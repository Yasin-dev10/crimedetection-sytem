from pathlib import Path
from docx import Document
from pypdf import PdfReader
import sys


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "3 chapters"


def extract_docx(path: Path) -> None:
    doc = Document(path)
    print(f"\n===== DOCX: {path.name} =====")
    for index, paragraph in enumerate(doc.paragraphs):
        text = " ".join(paragraph.text.split())
        if text:
            print(f"[P{index:03d} | {paragraph.style.name}] {text}")
    for table_index, table in enumerate(doc.tables):
        print(f"\n[TABLE {table_index}]")
        for row in table.rows:
            print(" | ".join(" ".join(cell.text.split()) for cell in row.cells))


def extract_pdf(path: Path) -> None:
    reader = PdfReader(path)
    print(f"\n===== PDF: {path.name} ({len(reader.pages)} pages) =====")
    for page_index, page in enumerate(reader.pages, start=1):
        text = " ".join((page.extract_text() or "").split())
        print(f"[PAGE {page_index:03d}] {text}")


def save_plain_text(path: Path) -> None:
    reader = PdfReader(path)
    pages = []
    for page_index, page in enumerate(reader.pages, start=1):
        pages.append(f"\n===== PAGE {page_index} =====\n{page.extract_text() or ''}")
    target = path.with_name(path.stem + "_extracted.txt")
    target.write_text("\n".join(pages), encoding="utf-8")
    print(target)


def save_docx_text(path: Path) -> None:
    doc = Document(path)
    lines = []
    for index, paragraph in enumerate(doc.paragraphs):
        text = " ".join(paragraph.text.split())
        if text:
            lines.append(f"[P{index:03d} | {paragraph.style.name}] {text}")
    for table_index, table in enumerate(doc.tables):
        lines.append(f"\n[TABLE {table_index}]")
        for row in table.rows:
            lines.append(" | ".join(" ".join(cell.text.split()) for cell in row.cells))
    target = path.with_name(path.stem + "_extracted.txt")
    target.write_text("\n".join(lines), encoding="utf-8")
    print(target)


if __name__ == "__main__":
    if len(sys.argv) > 1 and sys.argv[1] == "--save-pdf-text":
        save_plain_text(SOURCE_DIR / "Artifical Intelligence Based FYP.pdf")
        save_plain_text(
            SOURCE_DIR
            / "DEVELOPMENT OF AN EARLY WARNING AND PREDICTION SYSTEM FOR LOCUST OUTBREAKS USING MACHINE LEARNING  TE.pdf"
        )
        raise SystemExit
    if len(sys.argv) > 1 and sys.argv[1] == "--save-docx-text":
        save_docx_text(SOURCE_DIR / "chapter-one , twoo & three_backup.docx")
        raise SystemExit
    extract_docx(SOURCE_DIR / "chapter-one , twoo & three_backup.docx")
    extract_pdf(SOURCE_DIR / "Artifical Intelligence Based FYP.pdf")
    extract_pdf(
        SOURCE_DIR
        / "DEVELOPMENT OF AN EARLY WARNING AND PREDICTION SYSTEM FOR LOCUST OUTBREAKS USING MACHINE LEARNING  TE.pdf"
    )
